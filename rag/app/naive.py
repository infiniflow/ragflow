#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import logging
import re
import os
from io import BytesIO
from timeit import default_timer as timer
from docx import Document
from docx.image.exceptions import InvalidImageStreamError, UnexpectedEndOfFileError, UnrecognizedImageError
from docx.opc.pkgreader import _SerializedRelationships, _SerializedRelationship
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from docx.opc.oxml import parse_xml
from markdown import markdown
from PIL import Image

from common.constants import LLMType
from api.db.services.llm_service import LLMBundle
from rag.utils.file_utils import extract_embed_file, extract_links_from_pdf, extract_links_from_docx, extract_html
from deepdoc.parser import DocxParser, ExcelParser, HtmlParser, JsonParser,PdfParser, TxtParser ,MarkdownParser, MarkdownElementExtractor
from deepdoc.parser.figure_parser import adv_vision_figure_parser_wrapper, vision_figure_parser_pdf_wrapper
from deepdoc.parser.pdf_parser import PlainParser, VisionParser
from deepdoc.parser.docling_parser import DoclingParser
from deepdoc.parser.tcadp_parser import TCADPParser
from common.parser_config_utils import normalize_layout_recognizer
from rag.nlp import (
    concat_img,
    find_codec,
    naive_merge,
    naive_merge_with_images,
    adv_naive_merge,
    rag_tokenizer,
    tokenize_chunks,
    adv_tokenize_chunks,
    tokenize_table,
    append_context2table_image4pdf,
    tokenize_chunks_with_images,
)  # noqa: F401


def by_deepdoc(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, pdf_cls=None, **kwargs):
    callback = callback
    binary = binary
    pdf_parser = pdf_cls() if pdf_cls else Pdf()
    sections, tables = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page, callback=callback)

    tables = vision_figure_parser_pdf_wrapper(
        tbls=tables,
        sections=sections,
        callback=callback,
        **kwargs,
    )
    return sections, tables, pdf_parser


def by_mineru(
    filename,
    binary=None,
    from_page=0,
    to_page=100000,
    lang="Chinese",
    callback=None,
    pdf_cls=None,
    parse_method: str = "raw",
    mineru_llm_name: str | None = None,
    tenant_id: str | None = None,
    **kwargs,
):
    pdf_parser = None
    if tenant_id:
        if not mineru_llm_name:
            try:
                from api.db.services.tenant_llm_service import TenantLLMService

                env_name = TenantLLMService.ensure_mineru_from_env(tenant_id)
                candidates = TenantLLMService.query(tenant_id=tenant_id, llm_factory="MinerU", model_type=LLMType.OCR)
                if candidates:
                    mineru_llm_name = candidates[0].llm_name
                elif env_name:
                    mineru_llm_name = env_name
            except Exception as e:  # best-effort fallback
                logging.warning(f"fallback to env mineru: {e}")

        if mineru_llm_name:
            try:
                ocr_model = LLMBundle(tenant_id=tenant_id, llm_type=LLMType.OCR, llm_name=mineru_llm_name, lang=lang)
                pdf_parser = ocr_model.mdl
                sections, tables = pdf_parser.parse_pdf(
                    filepath=filename,
                    binary=binary,
                    callback=callback,
                    parse_method=parse_method,
                    lang=lang,
                    **kwargs,
                )
                return sections, tables, pdf_parser
            except Exception as e:
                logging.error(f"Failed to parse pdf via LLMBundle MinerU ({mineru_llm_name}): {e}")

    if callback:
        callback(-1, "MinerU not found.")
    return None, None, None


def by_docling(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, pdf_cls=None, **kwargs):
    pdf_parser = DoclingParser()
    parse_method = kwargs.get("parse_method", "raw")

    if not pdf_parser.check_installation():
        callback(-1, "Docling not found.")
        return None, None, pdf_parser

    sections, tables = pdf_parser.parse_pdf(
        filepath=filename,
        binary=binary,
        callback=callback,
        output_dir=os.environ.get("MINERU_OUTPUT_DIR", ""),
        delete_output=bool(int(os.environ.get("MINERU_DELETE_OUTPUT", 1))),
        parse_method=parse_method,
    )
    return sections, tables, pdf_parser


def by_tcadp(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, pdf_cls=None, **kwargs):
    tcadp_parser = TCADPParser()

    if not tcadp_parser.check_installation():
        callback(-1, "TCADP parser not available. Please check Tencent Cloud API configuration.")
        return None, None, tcadp_parser

    sections, tables = tcadp_parser.parse_pdf(filepath=filename, binary=binary, callback=callback, output_dir=os.environ.get("TCADP_OUTPUT_DIR", ""), file_type="PDF")
    return sections, tables, tcadp_parser


def by_paddleocr(
    filename,
    binary=None,
    from_page=0,
    to_page=100000,
    lang="Chinese",
    callback=None,
    pdf_cls=None,
    parse_method: str = "raw",
    paddleocr_llm_name: str | None = None,
    tenant_id: str | None = None,
    **kwargs,
):
    pdf_parser = None
    if tenant_id:
        if not paddleocr_llm_name:
            try:
                from api.db.services.tenant_llm_service import TenantLLMService

                env_name = TenantLLMService.ensure_paddleocr_from_env(tenant_id)
                candidates = TenantLLMService.query(tenant_id=tenant_id, llm_factory="PaddleOCR", model_type=LLMType.OCR)
                if candidates:
                    paddleocr_llm_name = candidates[0].llm_name
                elif env_name:
                    paddleocr_llm_name = env_name
            except Exception as e:  # best-effort fallback
                logging.warning(f"fallback to env paddleocr: {e}")

        if paddleocr_llm_name:
            try:
                ocr_model = LLMBundle(tenant_id=tenant_id, llm_type=LLMType.OCR, llm_name=paddleocr_llm_name, lang=lang)
                pdf_parser = ocr_model.mdl
                sections, tables = pdf_parser.parse_pdf(
                    filepath=filename,
                    binary=binary,
                    callback=callback,
                    parse_method=parse_method,
                    **kwargs,
                )
                return sections, tables, pdf_parser
            except Exception as e:
                logging.error(f"Failed to parse pdf via LLMBundle PaddleOCR ({paddleocr_llm_name}): {e}")

        return None, None, None

    if callback:
        callback(-1, "PaddleOCR not found.")
    return None, None, None


def by_plaintext(filename, binary=None, from_page=0, to_page=100000, callback=None, **kwargs):
    layout_recognizer = (kwargs.get("layout_recognizer") or "").strip()
    if (not layout_recognizer) or (layout_recognizer == "Plain Text"):
        pdf_parser = PlainParser()
    else:
        tenant_id = kwargs.get("tenant_id")
        if not tenant_id:
            raise ValueError("tenant_id is required when using vision layout recognizer")
        vision_model = LLMBundle(
            tenant_id,
            LLMType.IMAGE2TEXT,
            llm_name=layout_recognizer,
            lang=kwargs.get("lang", "Chinese"),
        )
        pdf_parser = VisionParser(vision_model=vision_model, **kwargs)

    sections, tables = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page, callback=callback)
    return sections, tables, pdf_parser


PARSERS = {
    "deepdoc": by_deepdoc,
    "mineru": by_mineru,
    "docling": by_docling,
    "tcadp": by_tcadp,
    "paddleocr": by_paddleocr,
    "plaintext": by_plaintext,  # default
}


class Docx(DocxParser):
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        imgs = paragraph._element.xpath(".//pic:pic")
        if not imgs:
            return None
        res_img = None
        for img in imgs:
            embed = img.xpath(".//a:blip/@r:embed")
            if not embed:
                continue
            embed = embed[0]
            try:
                related_part = document.part.related_parts[embed]
                image_blob = related_part.image.blob
            except UnrecognizedImageError:
                logging.info("Unrecognized image format. Skipping image.")
                continue
            except UnexpectedEndOfFileError:
                logging.info("EOF was unexpectedly encountered while reading an image stream. Skipping image.")
                continue
            except InvalidImageStreamError:
                logging.info("The recognized image stream appears to be corrupted. Skipping image.")
                continue
            except UnicodeDecodeError:
                logging.info("The recognized image stream appears to be corrupted. Skipping image.")
                continue
            except Exception as e:
                logging.warning(f"The recognized image stream appears to be corrupted. Skipping image, exception: {e}")
                continue
            try:
                image = Image.open(BytesIO(image_blob)).convert("RGB")
                if res_img is None:
                    res_img = image
                else:
                    res_img = concat_img(res_img, image)
            except Exception as e:
                logging.warning(f"Fail to open or concat images, exception: {e}")
                continue

        return res_img

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def __get_nearest_title(self, table_index, filename):
        """Get the hierarchical title structure before the table"""
        import re
        from docx.text.paragraph import Paragraph

        titles = []
        blocks = []

        # Get document name from filename parameter
        doc_name = re.sub(r"\.[a-zA-Z]+$", "", filename)
        if not doc_name:
            doc_name = "Untitled Document"

        # Collect all document blocks while maintaining document order
        try:
            # Iterate through all paragraphs and tables in document order
            for i, block in enumerate(self.doc._element.body):
                if block.tag.endswith("p"):  # Paragraph
                    p = Paragraph(block, self.doc)
                    blocks.append(("p", i, p))
                elif block.tag.endswith("tbl"):  # Table
                    blocks.append(("t", i, None))  # Table object will be retrieved later
        except Exception as e:
            logging.error(f"Error collecting blocks: {e}")
            return ""

        # Find the target table position
        target_table_pos = -1
        table_count = 0
        for i, (block_type, pos, _) in enumerate(blocks):
            if block_type == "t":
                if table_count == table_index:
                    target_table_pos = pos
                    break
                table_count += 1

        if target_table_pos == -1:
            return ""  # Target table not found

        # Find the nearest heading paragraph in reverse order
        nearest_title = None
        for i in range(len(blocks) - 1, -1, -1):
            block_type, pos, block = blocks[i]
            if pos >= target_table_pos:  # Skip blocks after the table
                continue

            if block_type != "p":
                continue

            if block.style and block.style.name and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                try:
                    level_match = re.search(r"(\d+)", block.style.name)
                    if level_match:
                        level = int(level_match.group(1))
                        if level <= 7:  # Support up to 7 heading levels
                            title_text = block.text.strip()
                            if title_text:  # Avoid empty titles
                                nearest_title = (level, title_text)
                                break
                except Exception as e:
                    logging.error(f"Error parsing heading level: {e}")

        if nearest_title:
            # Add current title
            titles.append(nearest_title)
            current_level = nearest_title[0]

            # Find all parent headings, allowing cross-level search
            while current_level > 1:
                found = False
                for i in range(len(blocks) - 1, -1, -1):
                    block_type, pos, block = blocks[i]
                    if pos >= target_table_pos:  # Skip blocks after the table
                        continue

                    if block_type != "p":
                        continue

                    if block.style and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                        try:
                            level_match = re.search(r"(\d+)", block.style.name)
                            if level_match:
                                level = int(level_match.group(1))
                                # Find any heading with a higher level
                                if level < current_level:
                                    title_text = block.text.strip()
                                    if title_text:  # Avoid empty titles
                                        titles.append((level, title_text))
                                        current_level = level
                                        found = True
                                        break
                        except Exception as e:
                            logging.error(f"Error parsing parent heading: {e}")

                if not found:  # Break if no parent heading is found
                    break

            # Sort by level (ascending, from highest to lowest)
            titles.sort(key=lambda x: x[0])
            # Organize titles (from highest to lowest)
            hierarchy = [doc_name] + [t[1] for t in titles]
            return " > ".join(hierarchy)

        return ""

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        last_image = None
        table_idx = 0

        def flush_last_image():
            nonlocal last_image, lines
            if last_image is not None:
                lines.append({"text": "", "image": last_image, "table": None, "style": "Image"})
                last_image = None

        for block in self.doc._element.body:
            if pn > to_page:
                break

            if block.tag.endswith("p"):
                p = Paragraph(block, self.doc)

                if from_page <= pn < to_page:
                    text = p.text.strip()
                    style_name = p.style.name if p.style else ""

                    if text:
                        if style_name == "Caption":
                            former_image = None

                            if lines and lines[-1].get("image") and lines[-1].get("style") != "Caption":
                                former_image = lines[-1].get("image")
                                lines.pop()

                            elif last_image is not None:
                                former_image = last_image
                                last_image = None

                            lines.append(
                                {
                                    "text": self.__clean(text),
                                    "image": former_image if former_image else None,
                                    "table": None,
                                }
                            )

                        else:
                            flush_last_image()
                            lines.append(
                                {
                                    "text": self.__clean(text),
                                    "image": None,
                                    "table": None,
                                }
                            )

                            current_image = self.get_picture(self.doc, p)
                            if current_image is not None:
                                lines.append(
                                    {
                                        "text": "",
                                        "image": current_image,
                                        "table": None,
                                    }
                                )

                    else:
                        current_image = self.get_picture(self.doc, p)
                        if current_image is not None:
                            last_image = current_image

                for run in p.runs:
                    xml = run._element.xml
                    if "lastRenderedPageBreak" in xml:
                        pn += 1
                        continue
                    if "w:br" in xml and 'type="page"' in xml:
                        pn += 1

            elif block.tag.endswith("tbl"):
                if pn < from_page or pn > to_page:
                    table_idx += 1
                    continue

                flush_last_image()
                tb = DocxTable(block, self.doc)
                title = self.__get_nearest_title(table_idx, filename)
                html = "<table>"
                if title:
                    html += f"<caption>Table Location: {title}</caption>"
                for r in tb.rows:
                    html += "<tr>"
                    col_idx = 0
                    try:
                        while col_idx < len(r.cells):
                            span = 1
                            c = r.cells[col_idx]
                            for j in range(col_idx + 1, len(r.cells)):
                                if c.text == r.cells[j].text:
                                    span += 1
                                    col_idx = j
                                else:
                                    break
                            col_idx += 1
                            html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                    except Exception as e:
                        logging.warning(f"Error parsing table, ignore: {e}")
                    html += "</tr>"
                html += "</table>"
                lines.append({"text": "", "image": None, "table": html})
                table_idx += 1

        flush_last_image()
        new_line = [(line.get("text"), line.get("image"), line.get("table")) for line in lines]

        return new_line

    def to_markdown(self, filename=None, binary=None, inline_images: bool = True):
        """
        This function uses mammoth, licensed under the BSD 2-Clause License.
        """

        import base64
        import uuid

        import mammoth
        from markdownify import markdownify

        docx_file = BytesIO(binary) if binary else open(filename, "rb")

        def _convert_image_to_base64(image):
            try:
                with image.open() as image_file:
                    image_bytes = image_file.read()
                encoded = base64.b64encode(image_bytes).decode("utf-8")
                base64_url = f"data:{image.content_type};base64,{encoded}"

                alt_name = "image"
                alt_name = f"img_{uuid.uuid4().hex[:8]}"

                return {"src": base64_url, "alt": alt_name}
            except Exception as e:
                logging.warning(f"Failed to convert image to base64: {e}")
                return {"src": "", "alt": "image"}

        try:
            if inline_images:
                result = mammoth.convert_to_html(docx_file, convert_image=mammoth.images.img_element(_convert_image_to_base64))
            else:
                result = mammoth.convert_to_html(docx_file)

            html = result.value

            markdown_text = markdownify(html)
            return markdown_text

        finally:
            if not binary:
                docx_file.close()


class Pdf(PdfParser):
    def __init__(self):
        super().__init__()

    def __call__(self, filename, binary=None, from_page=0, to_page=100000, zoomin=3, callback=None, separate_tables_figures=False):
        start = timer()
        first_start = start
        callback(msg="OCR started")
        self.__images__(filename if not binary else binary, zoomin, from_page, to_page, callback)
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        logging.info("OCR({}~{}): {:.2f}s".format(from_page, to_page, timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.63, "Layout analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._table_transformer_job(zoomin)
        callback(0.65, "Table analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._text_merge(zoomin=zoomin)
        callback(0.67, "Text merged ({:.2f}s)".format(timer() - start))

        if separate_tables_figures:
            tbls, figures = self._extract_table_figure(True, zoomin, True, True, True)
            self._concat_downward()
            logging.info("layouts cost: {}s".format(timer() - first_start))
            return [(b["text"], self._line_tag(b, zoomin)) for b in self.boxes], tbls, figures
        else:
            tbls = self._extract_table_figure(True, zoomin, True, True)
            self._naive_vertical_merge()
            self._concat_downward()
            # self._final_reading_order_merge()
            # self._filter_forpages()
            logging.info("layouts cost: {}s".format(timer() - first_start))
            return [(b["text"], self._line_tag(b, zoomin)) for b in self.boxes], tbls


class Markdown(MarkdownParser):
    def md_to_html(self, sections):
        if not sections:
            return []
        if isinstance(sections, type("")):
            text = sections
        elif isinstance(sections[0], type("")):
            text = sections[0]
        else:
            return []

        from bs4 import BeautifulSoup

        html_content = markdown(text)
        soup = BeautifulSoup(html_content, "html.parser")
        return soup

    def get_hyperlink_urls(self, soup):
        if soup:
            return set([a.get("href") for a in soup.find_all("a") if a.get("href")])
        return []
    

    def extract_image_urls_with_lines(self, segments):
        """
        Input: segments: List[Tuple[str, Any]]
        Output: List[Tuple[str, Any]]
        - Iterate original segments in order.
        - For non-text: append as-is.
        - For ("text", content): split by markdown/html image syntaxes, emitting:
                ("text", before), ("image", url), ("text", after) ...
        preserving the exact order.
        - Image segment payload is the extracted URL (string).
        """
        import re

        # Markdown image: ![alt](url)
        md_img_re = re.compile(r"!\[[^\]]*\]\(([^)\s]+)\)")

        # HTML image: <img ... src="url" ...>
        html_img_tag_re = re.compile(
            r"<img\b[^>]*\bsrc\s*=\s*(['\"])([^'\"\s>]+)\1[^>]*>",
            re.IGNORECASE | re.DOTALL,
        )

        out = []

        def split_text_keep_url(text: str):
            if text is None:
                text = ""

            matches = []

            # Markdown matches
            for m in md_img_re.finditer(text):
                matches.append((m.start(), m.end(), m.group(1)))  # url

            # HTML matches
            for m in html_img_tag_re.finditer(text):
                matches.append((m.start(), m.end(), m.group(2)))  # url

            if not matches:
                return [("text", text)]

            matches.sort(key=lambda x: x[0])

            res = []
            cursor = 0
            for s, e, url in matches:
                if s < cursor:
                    continue

                if s > cursor:
                    res.append(("text", text[cursor:s]))

                res.append(("image", url))
                cursor = e

            if cursor < len(text):
                res.append(("text", text[cursor:]))

            return res

        for seg_type, seg_content in segments:
            if seg_type != "text":
                out.append((seg_type, seg_content))
                continue

            out.extend(split_text_keep_url(seg_content))

        return out


    def load_images_from_urls(self, url, cache=None):
        import requests
        from pathlib import Path

        cache = cache or {}
        if url in cache:
            if cache[url]:
                return cache[url], cache
        img_obj = None
        try:
            #TODO： headers={'User-Agent': 'MyCustomAgent/1.0'}
            if url.startswith(("http://", "https://")):
                response = requests.get(url, stream=True, timeout=30)
                if response.status_code == 200 and response.headers.get("Content-Type", "").startswith("image/"):
                    img_obj = Image.open(BytesIO(response.content)).convert("RGB")
            else:
                local_path = Path(url)
                if local_path.exists():
                    img_obj = Image.open(url).convert("RGB")
                else:
                    logging.warning(f"Local image file not found: {url}")
        except Exception as e:
            logging.error(f"Failed to download/open image from {url}: {e}")
        cache[url] = img_obj
        return img_obj, cache


    def __call__(self, filename, binary=None, delimiter=None):
        if binary:
            encoding = find_codec(binary)
            txt = binary.decode(encoding, errors="ignore")
        else:
            with open(filename, "r") as f:
                txt = f.read()

        # 1) Raw markdown → segments (text / table)
        segments = self.extract_table_segments(f"{txt}\n")

        # 2) Split images from text → segments (text / image / table)
        segments = self.extract_image_urls_with_lines(segments)

        # 3) Split text into markdown blocks (text only; others passthrough)
        extractor = MarkdownElementExtractor(segments)
        segments = extractor.extract_elements(delimiter, include_meta=True)

        sections = []
        image_cache = {}
        for seg in segments:
            seg_type, seg_content = seg
            if seg_type == "text":
                sections.append((seg_content, "", ""))
            elif seg_type == "image":
                imgs, image_cache = self.load_images_from_urls(seg_content, image_cache)
                sections.append(("", imgs, ""))
            elif seg_type == "table":
                sections.append(("","",markdown(seg_content, extensions=["markdown.extensions.tables"])))

        return sections
    

def load_from_xml_v2(baseURI, rels_item_xml):
    """
    Return |_SerializedRelationships| instance loaded with the
    relationships contained in *rels_item_xml*. Returns an empty
    collection if *rels_item_xml* is |None|.
    """
    srels = _SerializedRelationships()
    if rels_item_xml is not None:
        rels_elm = parse_xml(rels_item_xml)
        for rel_elm in rels_elm.Relationship_lst:
            if rel_elm.target_ref in ("../NULL", "NULL") or rel_elm.target_ref.startswith("#"):
                continue
            srels._srels.append(_SerializedRelationship(baseURI, rel_elm))
    return srels


def chunk(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, **kwargs):
    """
    Supported file formats are docx, pdf, excel, txt.
    This method apply the naive ways to chunk files.
    Successive text will be sliced into pieces using 'delimiter'.
    Next, these successive pieces are merge into chunks whose token number is no more than 'Max token number'.
    """
    urls = set()
    url_res = []

    is_english = lang.lower() == "english"  # is_english(cks)
    parser_config = kwargs.get(
        "parser_config",
        {
            "chunk_token_num": 512,
            "delimiter": "\n!?。；！？",
            "layout_recognize": "DeepDOC",
            "analyze_hyperlink": True,
        },
    )

    # Normalize children delimiters and preserve custom tokens wrapped by backticks.
    raw_child_deli = parser_config.get("children_delimiter") or ""
    child_deli = raw_child_deli.encode("utf-8").decode("unicode_escape").encode("latin1").decode("utf-8")
    cust_child_deli = re.findall(r"`([^`]+)`", child_deli)
    child_deli = "|".join(re.sub(r"`([^`]+)`", "", child_deli))
    if cust_child_deli:
        cust_child_deli = sorted(set(cust_child_deli), key=lambda x: -len(x))
        cust_child_deli = "|".join(re.escape(t) for t in cust_child_deli if t)
        child_deli += cust_child_deli

    is_markdown = False
    table_context_size = max(0, int(parser_config.get("table_context_size", 0) or 0))
    image_context_size = max(0, int(parser_config.get("image_context_size", 0) or 0))

    overlapped_percent = int(parser_config.get("overlapped_percent", 0))
    overlapped_percent = max(0, min(overlapped_percent, 90))

    # Document-level tokens for title-aware chunking.
    doc = {"docnm_kwd": filename, "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))}
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    res = []
    pdf_parser = None
    section_images = None

    is_root = kwargs.get("is_root", True)
    embed_res = []
    if is_root:
        # Only extract embedded files once, then merge their chunking results.
        embeds = []
        if binary is not None:
            embeds = extract_embed_file(binary)
        else:
            raise Exception("Embedding extraction from file path is not supported.")

        # Recursively chunk each embedded file and collect results
        for embed_filename, embed_bytes in embeds:
            try:
                sub_res = chunk(embed_filename, binary=embed_bytes, lang=lang, callback=callback, is_root=False, **kwargs) or []
                embed_res.extend(sub_res)
            except Exception as e:
                error_msg = f"Failed to chunk embed {embed_filename}: {e}"
                logging.error(error_msg)
                if callback:
                    callback(0.05, error_msg)
                continue

    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        if parser_config.get("analyze_hyperlink", False) and is_root:
            urls = extract_links_from_docx(binary)
            for index, url in enumerate(urls):
                html_bytes, metadata = extract_html(url)
                if not html_bytes:
                    continue
                try:
                    sub_url_res = chunk(url, html_bytes, callback=callback, lang=lang, is_root=False, **kwargs)
                except Exception as e:
                    logging.info(f"Failed to chunk url in registered file type {url}: {e}")
                    sub_url_res = chunk(f"{index}.html", html_bytes, callback=callback, lang=lang, is_root=False, **kwargs)
                url_res.extend(sub_url_res)

        # fix "There is no item named 'word/NULL' in the archive", referring to https://github.com/python-openxml/python-docx/issues/1105#issuecomment-1298075246
        _SerializedRelationships.load_from_xml = load_from_xml_v2

        # sections = (text, image, tables)
        sections = Docx()(filename, binary)

        # chunks list[dict]
        # images list - index of image chunk in chunks
        chunks, images = adv_naive_merge(sections, int(parser_config.get("chunk_token_num", 128)), parser_config.get("delimiter", "\n!?。；！？"), table_context_size, image_context_size)

        adv_vision_figure_parser_wrapper(chunks=chunks, idx_lst=images, callback=callback, **kwargs)

        callback(0.8, "Finish parsing.")
        st = timer()

        res.extend(adv_tokenize_chunks(chunks, doc, is_english, child_delimiters_pattern=child_deli))
        logging.info("naive_merge({}): {}".format(filename, timer() - st))
        res.extend(embed_res)
        res.extend(url_res)
        return res

    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        layout_recognizer, parser_model_name = normalize_layout_recognizer(parser_config.get("layout_recognize", "DeepDOC"))

        if parser_config.get("analyze_hyperlink", False) and is_root:
            urls = extract_links_from_pdf(binary)

        if isinstance(layout_recognizer, bool):
            layout_recognizer = "DeepDOC" if layout_recognizer else "Plain Text"

        name = layout_recognizer.strip().lower()
        parser = PARSERS.get(name, by_plaintext)
        callback(0.1, "Start to parse.")

        sections, tables, pdf_parser = parser(
            filename=filename,
            binary=binary,
            from_page=from_page,
            to_page=to_page,
            lang=lang,
            callback=callback,
            layout_recognizer=layout_recognizer,
            mineru_llm_name=parser_model_name,
            paddleocr_llm_name=parser_model_name,
            **kwargs,
        )

        if not sections and not tables:
            return []

        if table_context_size or image_context_size:
            tables = append_context2table_image4pdf(sections, tables, image_context_size)

        if name in ["tcadp", "docling", "mineru", "paddleocr"]:
            parser_config["chunk_token_num"] = 0

        res = tokenize_table(tables, doc, is_english)
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(csv|xlsx?)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")

        # Check if tcadp_parser is selected for spreadsheet files
        layout_recognizer = parser_config.get("layout_recognize", "DeepDOC")
        if layout_recognizer == "TCADP Parser":
            table_result_type = parser_config.get("table_result_type", "1")
            markdown_image_response_type = parser_config.get("markdown_image_response_type", "1")
            tcadp_parser = TCADPParser(table_result_type=table_result_type, markdown_image_response_type=markdown_image_response_type)
            if not tcadp_parser.check_installation():
                callback(-1, "TCADP parser not available. Please check Tencent Cloud API configuration.")
                return res

            # Determine file type based on extension
            file_type = "XLSX" if re.search(r"\.xlsx?$", filename, re.IGNORECASE) else "CSV"

            sections, tables = tcadp_parser.parse_pdf(filepath=filename, binary=binary, callback=callback, output_dir=os.environ.get("TCADP_OUTPUT_DIR", ""), file_type=file_type)
            parser_config["chunk_token_num"] = 0
            res = tokenize_table(tables, doc, is_english)
            callback(0.8, "Finish parsing.")
        else:
            # Default DeepDOC parser
            excel_parser = ExcelParser()
            if parser_config.get("html4excel"):
                sections = [(_, "") for _ in excel_parser.html(binary, 12) if _]
                parser_config["chunk_token_num"] = 0
            else:
                sections = [(_, "") for _ in excel_parser(binary) if _]

    elif re.search(r"\.(txt|py|js|java|c|cpp|h|php|go|ts|sh|cs|kt|sql)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = TxtParser()(filename, binary, parser_config.get("chunk_token_num", 128), parser_config.get("delimiter", "\n!?;。；！？"))
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(md|markdown|mdx)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        markdown_parser = Markdown(int(parser_config.get("chunk_token_num", 128)))
        sections = markdown_parser(
            filename,
            binary,
            delimiter=parser_config.get("delimiter", "\n!?;。；！？"),
        )

        is_markdown = True

        chunks, images = adv_naive_merge(sections, int(parser_config.get("chunk_token_num", 128)), "\n。；！？", table_context_size, image_context_size)
        
        adv_vision_figure_parser_wrapper(chunks=chunks, idx_lst=images, callback=callback, **kwargs)

        if parser_config.get("hyperlink_urls", False) and is_root:
            for idx, (section_text, _) in enumerate(sections):
                soup = markdown_parser.md_to_html(section_text)
                hyperlink_urls = markdown_parser.get_hyperlink_urls(soup)
                urls.update(hyperlink_urls)

        callback(0.8, "Finish parsing.")
        res.extend(adv_tokenize_chunks(chunks, doc, is_english, child_delimiters_pattern=child_deli))

    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        chunk_token_num = int(parser_config.get("chunk_token_num", 128))
        sections = HtmlParser()(filename, binary, chunk_token_num)
        sections = [(_, "") for _ in sections if _]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(json|jsonl|ldjson)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        chunk_token_num = int(parser_config.get("chunk_token_num", 128))
        sections = JsonParser(chunk_token_num)(binary)
        sections = [(_, "") for _ in sections if _]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")

        try:
            from tika import parser as tika_parser
        except Exception as e:
            callback(0.8, f"tika not available: {e}. Unsupported .doc parsing.")
            logging.warning(f"tika not available: {e}. Unsupported .doc parsing for {filename}.")
            return []

        binary = BytesIO(binary)
        doc_parsed = tika_parser.from_buffer(binary)
        if doc_parsed.get("content", None) is not None:
            sections = doc_parsed["content"].split("\n")
            sections = [(_, "") for _ in sections if _]
            callback(0.8, "Finish parsing.")
        else:
            error_msg = f"tika.parser got empty content from {filename}."
            callback(0.8, error_msg)
            logging.warning(error_msg)
            return []
    else:
        raise NotImplementedError("file type not supported yet(pdf, xlsx, doc, docx, txt supported)")

    st = timer()
    if not is_markdown:
        if section_images:
            if all(image is None for image in section_images):
                section_images = None

        if section_images:
            chunks, images = naive_merge_with_images(sections, section_images, int(parser_config.get("chunk_token_num", 128)), parser_config.get("delimiter", "\n!?。；！？"), overlapped_percent)
            res.extend(tokenize_chunks_with_images(chunks, doc, is_english, images, child_delimiters_pattern=child_deli))
        else:
            chunks = naive_merge(sections, int(parser_config.get("chunk_token_num", 128)), parser_config.get("delimiter", "\n!?。；！？"), overlapped_percent)

            res.extend(tokenize_chunks(chunks, doc, is_english, pdf_parser, child_delimiters_pattern=child_deli))

    if urls and parser_config.get("analyze_hyperlink", False) and is_root:
        for index, url in enumerate(urls):
            html_bytes, metadata = extract_html(url)
            if not html_bytes:
                continue
            try:
                sub_url_res = chunk(url, html_bytes, callback=callback, lang=lang, is_root=False, **kwargs)
            except Exception as e:
                logging.info(f"Failed to chunk url in registered file type {url}: {e}")
                sub_url_res = chunk(f"{index}.html", html_bytes, callback=callback, lang=lang, is_root=False, **kwargs)
            url_res.extend(sub_url_res)

    logging.info("naive_merge({}): {}".format(filename, timer() - st))

    if embed_res:
        res.extend(embed_res)
    if url_res:
        res.extend(url_res)
    # if table_context_size or image_context_size:
    #    attach_media_context(res, table_context_size, image_context_size)
    return res


if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass

    chunk(sys.argv[1], from_page=0, to_page=10, callback=dummy)
