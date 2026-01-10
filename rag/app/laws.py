#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import logging
import re
from io import BytesIO
from docx import Document

from common.constants import ParserType
from deepdoc.parser.utils import get_text
from rag.nlp import bullets_category, remove_contents_table, make_colon_as_title, tokenize_chunks, docx_question_level, tree_merge
from rag.nlp import rag_tokenizer, Node
from deepdoc.parser import PdfParser, DocxParser, HtmlParser
from rag.app.naive import by_plaintext, PARSERS
from common.parser_config_utils import normalize_layout_recognizer


class Docx(DocxParser):
    def __init__(self):
        pass

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def old_call(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page and p.text.strip():
                lines.append(self.__clean(p.text))
            for run in p.runs:
                if "lastRenderedPageBreak" in run._element.xml:
                    pn += 1
                    continue
                if "w:br" in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        return [line for line in lines if line]

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        level_set = set()
        bull = bullets_category([p.text for p in self.doc.paragraphs])
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            question_level, p_text = docx_question_level(p, bull)
            if not p_text.strip("\n"):
                continue
            lines.append((question_level, p_text))
            level_set.add(question_level)
            for run in p.runs:
                if "lastRenderedPageBreak" in run._element.xml:
                    pn += 1
                    continue
                if "w:br" in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1

        sorted_levels = sorted(level_set)

        h2_level = sorted_levels[1] if len(sorted_levels) > 1 else 1
        h2_level = sorted_levels[-2] if h2_level == sorted_levels[-1] and len(sorted_levels) > 2 else h2_level

        root = Node(level=0, depth=h2_level, texts=[])
        root.build_tree(lines)

        return [element for element in root.get_tree() if element]

    def __str__(self) -> str:
        return f"""
            question:{self.question},
            answer:{self.answer},
            level:{self.level},
            childs:{self.childs}
        """


class Pdf(PdfParser):
    def __init__(self):
        self.model_speciess = ParserType.LAWS.value
        super().__init__()

    def __call__(self, filename, binary=None, from_page=0, to_page=100000, zoomin=3, callback=None):
        from timeit import default_timer as timer

        start = timer()
        callback(msg="OCR started")
        self.__images__(filename if not binary else binary, zoomin, from_page, to_page, callback)
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.67, "Layout analysis ({:.2f}s)".format(timer() - start))
        logging.debug("layouts: {}".format((timer() - start)))
        self._naive_vertical_merge()

        callback(0.8, "Text extraction ({:.2f}s)".format(timer() - start))

        return [(b["text"], self._line_tag(b, zoomin)) for b in self.boxes], None


def chunk(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, **kwargs):
    """
    Supported file formats are docx, pdf, txt.
    """
    parser_config = kwargs.get("parser_config", {"chunk_token_num": 512, "delimiter": "\n!?。；！？", "layout_recognize": "DeepDOC"})
    doc = {"docnm_kwd": filename, "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))}
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    pdf_parser = None
    sections = []
    # is it English
    eng = lang.lower() == "english"  # is_english(sections)

    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        chunks = Docx()(filename, binary)
        callback(0.7, "Finish parsing.")
        return tokenize_chunks(chunks, doc, eng, None)

    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        layout_recognizer, parser_model_name = normalize_layout_recognizer(parser_config.get("layout_recognize", "DeepDOC"))

        if isinstance(layout_recognizer, bool):
            layout_recognizer = "DeepDOC" if layout_recognizer else "Plain Text"

        name = layout_recognizer.strip().lower()
        parser = PARSERS.get(name, by_plaintext)
        callback(0.1, "Start to parse.")

        raw_sections, tables, pdf_parser = parser(
            filename=filename,
            binary=binary,
            from_page=from_page,
            to_page=to_page,
            lang=lang,
            callback=callback,
            pdf_cls=Pdf,
            layout_recognizer=layout_recognizer,
            mineru_llm_name=parser_model_name,
            paddleocr_llm_name=parser_model_name,
            **kwargs,
        )

        if not raw_sections and not tables:
            return []

        if name in ["tcadp", "docling", "mineru", "paddleocr"]:
            parser_config["chunk_token_num"] = 0

        for txt, poss in raw_sections:
            sections.append(txt + poss)

        callback(0.8, "Finish parsing.")
    elif re.search(r"\.(txt|md|markdown|mdx)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        sections = txt.split("\n")
        sections = [s for s in sections if s]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = HtmlParser()(filename, binary)
        sections = [s for s in sections if s]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        try:
            from tika import parser as tika_parser
        except Exception as e:
            callback(0.8, f"tika not available: {e}. Unsupported .doc parsing.")
            logging.warning(f"tika not available: {e}. Unsupported .doc parsing for {filename}.")
            return []

        binary = BytesIO(binary)
        doc_parsed = tika_parser.from_buffer(binary)
        if doc_parsed.get("content", None) is not None:
            sections = doc_parsed["content"].split("\n")
            sections = [s for s in sections if s]
            callback(0.8, "Finish parsing.")
        else:
            callback(0.8, f"tika.parser got empty content from {filename}.")
            logging.warning(f"tika.parser got empty content from {filename}.")
            return []
    else:
        raise NotImplementedError("file type not supported yet(doc, docx, pdf, txt supported)")

    # Remove 'Contents' part
    remove_contents_table(sections, eng)

    make_colon_as_title(sections)
    bull = bullets_category(sections)
    res = tree_merge(bull, sections, 2)

    if not res:
        callback(0.99, "No chunk parsed out.")

    return tokenize_chunks(res, doc, eng, pdf_parser)

    # chunks = hierarchical_merge(bull, sections, 5)
    #     return tokenize_chunks(["\n".join(ck)for ck in chunks], doc, eng, pdf_parser)


if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass

    chunk(sys.argv[1], callback=dummy)
