#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import logging
import copy
import re

from common.constants import ParserType
from io import BytesIO
from rag.nlp import rag_tokenizer, tokenize, tokenize_table, bullets_category, title_frequency, tokenize_chunks, docx_question_level, attach_media_context
from common.token_utils import num_tokens_from_string
from deepdoc.parser import PdfParser, DocxParser
from deepdoc.parser.figure_parser import vision_figure_parser_pdf_wrapper, vision_figure_parser_docx_wrapper
from docx import Document
from PIL import Image
from rag.app.naive import by_plaintext, PARSERS
from common.parser_config_utils import normalize_layout_recognizer


class Pdf(PdfParser):
    def __init__(self):
        self.model_speciess = ParserType.MANUAL.value
        super().__init__()

    def __call__(self, filename, binary=None, from_page=0, to_page=100000, zoomin=3, callback=None):
        from timeit import default_timer as timer

        start = timer()
        callback(msg="OCR started")
        self.__images__(filename if not binary else binary, zoomin, from_page, to_page, callback)
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        logging.debug("OCR: {}".format(timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.65, "Layout analysis ({:.2f}s)".format(timer() - start))
        logging.debug("layouts: {}".format(timer() - start))

        start = timer()
        self._table_transformer_job(zoomin)
        callback(0.67, "Table analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._text_merge()
        tbls = self._extract_table_figure(True, zoomin, True, True)
        self._concat_downward()
        self._filter_forpages()
        callback(0.68, "Text merged ({:.2f}s)".format(timer() - start))

        # clean mess
        for b in self.boxes:
            b["text"] = re.sub(r"([\t 　]|\u3000){2,}", " ", b["text"].strip())

        return [(b["text"], b.get("layoutno", ""), self.get_position(b, zoomin)) for i, b in enumerate(self.boxes)], tbls


class Docx(DocxParser):
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        img = paragraph._element.xpath(".//pic:pic")
        if not img:
            return None
        try:
            img = img[0]
            embed = img.xpath(".//a:blip/@r:embed")[0]
            related_part = document.part.related_parts[embed]
            image = related_part.image
            if image is not None:
                image = Image.open(BytesIO(image.blob))
                return image
            elif related_part.blob is not None:
                image = Image.open(BytesIO(related_part.blob))
                return image
            else:
                return None
        except Exception:
            return None

    def concat_img(self, img1, img2):
        if img1 and not img2:
            return img1
        if not img1 and img2:
            return img2
        if not img1 and not img2:
            return None
        width1, height1 = img1.size
        width2, height2 = img2.size

        new_width = max(width1, width2)
        new_height = height1 + height2
        new_image = Image.new("RGB", (new_width, new_height))

        new_image.paste(img1, (0, 0))
        new_image.paste(img2, (0, height1))

        return new_image

    def __call__(self, filename, binary=None, from_page=0, to_page=100000, callback=None):
        self.doc = Document(filename) if not binary else Document(BytesIO(binary))
        pn = 0
        last_answer, last_image = "", None
        question_stack, level_stack = [], []
        ti_list = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            question_level, p_text = 0, ""
            if from_page <= pn < to_page and p.text.strip():
                question_level, p_text = docx_question_level(p)
            if not question_level or question_level > 6:  # not a question
                last_answer = f"{last_answer}\n{p_text}"
                current_image = self.get_picture(self.doc, p)
                last_image = self.concat_img(last_image, current_image)
            else:  # is a question
                if last_answer or last_image:
                    sum_question = "\n".join(question_stack)
                    if sum_question:
                        ti_list.append((f"{sum_question}\n{last_answer}", last_image))
                    last_answer, last_image = "", None

                i = question_level
                while question_stack and i <= level_stack[-1]:
                    question_stack.pop()
                    level_stack.pop()
                question_stack.append(p_text)
                level_stack.append(question_level)
            for run in p.runs:
                if "lastRenderedPageBreak" in run._element.xml:
                    pn += 1
                    continue
                if "w:br" in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        if last_answer:
            sum_question = "\n".join(question_stack)
            if sum_question:
                ti_list.append((f"{sum_question}\n{last_answer}", last_image))

        tbls = []
        for tb in self.doc.tables:
            html = "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i + 1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                        else:
                            break
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return ti_list, tbls


def chunk(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, **kwargs):
    """
    Only pdf is supported.
    """
    parser_config = kwargs.get("parser_config", {"chunk_token_num": 512, "delimiter": "\n!?。；！？", "layout_recognize": "DeepDOC"})
    pdf_parser = None
    doc = {"docnm_kwd": filename}
    doc["title_tks"] = rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", doc["docnm_kwd"]))
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    # is it English
    eng = lang.lower() == "english"  # pdf_parser.is_english
    if re.search(r"\.pdf$", filename, re.IGNORECASE):
        layout_recognizer, parser_model_name = normalize_layout_recognizer(parser_config.get("layout_recognize", "DeepDOC"))

        if isinstance(layout_recognizer, bool):
            layout_recognizer = "DeepDOC" if layout_recognizer else "Plain Text"

        name = layout_recognizer.strip().lower()
        pdf_parser = PARSERS.get(name, by_plaintext)
        callback(0.1, "Start to parse.")

        kwargs.pop("parse_method", None)
        kwargs.pop("mineru_llm_name", None)
        sections, tbls, pdf_parser = pdf_parser(
            filename=filename,
            binary=binary,
            from_page=from_page,
            to_page=to_page,
            lang=lang,
            callback=callback,
            pdf_cls=Pdf,
            layout_recognizer=layout_recognizer,
            mineru_llm_name=parser_model_name,
            paddleocr_llm_name=parser_model_name,
            parse_method="manual",
            **kwargs,
        )

        def _normalize_section(section):
            # pad section to length 3: (txt, sec_id, poss)
            if len(section) == 1:
                section = (section[0], "", [])
            elif len(section) == 2:
                section = (section[0], "", section[1])
            elif len(section) != 3:
                raise ValueError(f"Unexpected section length: {len(section)} (value={section!r})")

            txt, layoutno, poss = section
            if isinstance(poss, str):
                poss = pdf_parser.extract_positions(poss)
                if poss:
                    first = poss[0]  # tuple: ([pn], x1, x2, y1, y2)
                    pn = first[0]
                    if isinstance(pn, list) and pn:
                        pn = pn[0]  # [pn] -> pn
                        poss[0] = (pn, *first[1:])

            return (txt, layoutno, poss)

        sections = [_normalize_section(sec) for sec in sections]

        if not sections and not tbls:
            return []

        if name in ["tcadp", "docling", "mineru", "paddleocr"]:
            parser_config["chunk_token_num"] = 0

        callback(0.8, "Finish parsing.")

        if len(sections) > 0 and len(pdf_parser.outlines) / len(sections) > 0.03:
            max_lvl = max([lvl for _, lvl in pdf_parser.outlines])
            most_level = max(0, max_lvl - 1)
            levels = []
            for txt, _, _ in sections:
                for t, lvl in pdf_parser.outlines:
                    tks = set([t[i] + t[i + 1] for i in range(len(t) - 1)])
                    tks_ = set([txt[i] + txt[i + 1] for i in range(min(len(t), len(txt) - 1))])
                    if len(set(tks & tks_)) / max([len(tks), len(tks_), 1]) > 0.8:
                        levels.append(lvl)
                        break
                else:
                    levels.append(max_lvl + 1)

        else:
            bull = bullets_category([txt for txt, _, _ in sections])
            most_level, levels = title_frequency(bull, [(txt, lvl) for txt, lvl, _ in sections])

        assert len(sections) == len(levels)
        sec_ids = []
        sid = 0
        for i, lvl in enumerate(levels):
            if lvl <= most_level and i > 0 and lvl != levels[i - 1]:
                sid += 1
            sec_ids.append(sid)

        sections = [(txt, sec_ids[i], poss) for i, (txt, _, poss) in enumerate(sections)]
        for (img, rows), poss in tbls:
            if not rows:
                continue
            sections.append((rows if isinstance(rows, str) else rows[0], -1, [(p[0] + 1 - from_page, p[1], p[2], p[3], p[4]) for p in poss]))

        def tag(pn, left, right, top, bottom):
            if pn + left + right + top + bottom == 0:
                return ""
            return "@@{}\t{:.1f}\t{:.1f}\t{:.1f}\t{:.1f}##".format(pn, left, right, top, bottom)

        chunks = []
        last_sid = -2
        tk_cnt = 0
        for txt, sec_id, poss in sorted(sections, key=lambda x: (x[-1][0][0], x[-1][0][3], x[-1][0][1])):
            poss = "\t".join([tag(*pos) for pos in poss])
            if tk_cnt < 32 or (tk_cnt < 1024 and (sec_id == last_sid or sec_id == -1)):
                if chunks:
                    chunks[-1] += "\n" + txt + poss
                    tk_cnt += num_tokens_from_string(txt)
                    continue
            chunks.append(txt + poss)
            tk_cnt = num_tokens_from_string(txt)
            if sec_id > -1:
                last_sid = sec_id
        tbls = vision_figure_parser_pdf_wrapper(
            tbls=tbls,
            sections=sections,
            callback=callback,
            **kwargs,
        )
        res = tokenize_table(tbls, doc, eng)
        res.extend(tokenize_chunks(chunks, doc, eng, pdf_parser))
        table_ctx = max(0, int(parser_config.get("table_context_size", 0) or 0))
        image_ctx = max(0, int(parser_config.get("image_context_size", 0) or 0))
        if table_ctx or image_ctx:
            attach_media_context(res, table_ctx, image_ctx)
        return res

    elif re.search(r"\.docx?$", filename, re.IGNORECASE):
        docx_parser = Docx()
        ti_list, tbls = docx_parser(filename, binary, from_page=0, to_page=10000, callback=callback)
        tbls = vision_figure_parser_docx_wrapper(sections=ti_list, tbls=tbls, callback=callback, **kwargs)
        res = tokenize_table(tbls, doc, eng)
        for text, image in ti_list:
            d = copy.deepcopy(doc)
            if image:
                d["image"] = image
                d["doc_type_kwd"] = "image"
            tokenize(d, text, eng)
            res.append(d)
        table_ctx = max(0, int(parser_config.get("table_context_size", 0) or 0))
        image_ctx = max(0, int(parser_config.get("image_context_size", 0) or 0))
        if table_ctx or image_ctx:
            attach_media_context(res, table_ctx, image_ctx)
        return res
    else:
        raise NotImplementedError("file type not supported yet(pdf and docx supported)")


if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass

    chunk(sys.argv[1], callback=dummy)
