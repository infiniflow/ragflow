#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import logging
import re
import csv
from copy import deepcopy
from io import BytesIO
from timeit import default_timer as timer
from openpyxl import load_workbook

from deepdoc.parser.utils import get_text
from rag.nlp import is_english, random_choices, qbullets_category, add_positions, has_qbullet, docx_question_level
from rag.nlp import rag_tokenizer, tokenize_table, concat_img
from deepdoc.parser import PdfParser, ExcelParser, DocxParser
from docx import Document
from PIL import Image
from markdown import markdown


class Excel(ExcelParser):
    def __call__(self, fnm, binary=None, callback=None):
        if not binary:
            wb = load_workbook(fnm)
        else:
            wb = load_workbook(BytesIO(binary))
        total = 0
        for sheetname in wb.sheetnames:
            total += len(list(wb[sheetname].rows))

        res, fails = [], []
        for sheetname in wb.sheetnames:
            ws = wb[sheetname]
            rows = list(ws.rows)
            for i, r in enumerate(rows):
                q, a = "", ""
                for cell in r:
                    if not cell.value:
                        continue
                    if not q:
                        q = str(cell.value)
                    elif not a:
                        a = str(cell.value)
                    else:
                        break
                if q and a:
                    res.append((q, a))
                else:
                    fails.append(str(i + 1))
                if len(res) % 999 == 0:
                    callback(len(res) *
                             0.6 /
                             total, ("Extract pairs: {}".format(len(res)) +
                                     (f"{len(fails)} failure, line: %s..." %
                                      (",".join(fails[:3])) if fails else "")))

        callback(0.6, ("Extract pairs: {}. ".format(len(res)) + (
            f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))
        self.is_english = is_english(
            [rmPrefix(q) for q, _ in random_choices(res, k=30) if len(q) > 1])
        return res


class Pdf(PdfParser):
    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None):
        start = timer()
        callback(msg="OCR started")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        logging.debug("OCR({}~{}): {:.2f}s".format(from_page, to_page, timer() - start))
        start = timer()
        self._layouts_rec(zoomin, drop=False)
        callback(0.63, "Layout analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._table_transformer_job(zoomin)
        callback(0.65, "Table analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._text_merge()
        callback(0.67, "Text merged ({:.2f}s)".format(timer() - start))
        tbls = self._extract_table_figure(True, zoomin, True, True)
        #self._naive_vertical_merge()
        # self._concat_downward()
        #self._filter_forpages()
        logging.debug("layouts: {}".format(timer() - start))
        sections = [b["text"] for b in self.boxes]
        bull_x0_list = []
        q_bull, reg = qbullets_category(sections)
        if q_bull == -1:
            raise ValueError("Unable to recognize Q&A structure.")
        qai_list = []
        last_q, last_a, last_tag = '', '', ''
        last_index = -1
        last_box = {'text':''}
        last_bull = None
        def sort_key(element):
            tbls_pn = element[1][0][0]
            tbls_top = element[1][0][3]
            return tbls_pn, tbls_top
        tbls.sort(key=sort_key)
        tbl_index = 0
        last_pn, last_bottom = 0, 0
        tbl_pn, tbl_left, tbl_right, tbl_top, tbl_bottom, tbl_tag, tbl_text = 1, 0, 0, 0, 0, '@@0\t0\t0\t0\t0##', ''
        for box in self.boxes:
            section, line_tag = box['text'], self._line_tag(box, zoomin)
            has_bull, index = has_qbullet(reg, box, last_box, last_index, last_bull, bull_x0_list)
            last_box, last_index, last_bull = box, index, has_bull
            line_pn = float(line_tag.lstrip('@@').split('\t')[0])
            line_top = float(line_tag.rstrip('##').split('\t')[3])
            tbl_pn, tbl_left, tbl_right, tbl_top, tbl_bottom, tbl_tag, tbl_text = self.get_tbls_info(tbls, tbl_index)
            if not has_bull:  # No question bullet
                if not last_q:
                    if tbl_pn < line_pn or (tbl_pn == line_pn and tbl_top <= line_top):    # image passed
                        tbl_index += 1
                    continue
                else:
                    sum_tag = line_tag
                    sum_section = section
                    while ((tbl_pn == last_pn and tbl_top>= last_bottom) or (tbl_pn > last_pn)) \
                        and ((tbl_pn == line_pn and tbl_top <= line_top) or (tbl_pn < line_pn)):    # add image at the middle of current answer
                        sum_tag = f'{tbl_tag}{sum_tag}'
                        sum_section = f'{tbl_text}{sum_section}'
                        tbl_index += 1
                        tbl_pn, tbl_left, tbl_right, tbl_top, tbl_bottom, tbl_tag, tbl_text = self.get_tbls_info(tbls, tbl_index)
                    last_a = f'{last_a}{sum_section}'
                    last_tag = f'{last_tag}{sum_tag}'
            else:
                if last_q:
                    while ((tbl_pn == last_pn and tbl_top>= last_bottom) or (tbl_pn > last_pn)) \
                        and ((tbl_pn == line_pn and tbl_top <= line_top) or (tbl_pn < line_pn)):    # add image at the end of last answer
                        last_tag = f'{last_tag}{tbl_tag}'
                        last_a = f'{last_a}{tbl_text}'
                        tbl_index += 1
                        tbl_pn, tbl_left, tbl_right, tbl_top, tbl_bottom, tbl_tag, tbl_text = self.get_tbls_info(tbls, tbl_index)
                    image, poss = self.crop(last_tag, need_position=True)
                    qai_list.append((last_q, last_a, image, poss))
                    last_q, last_a, last_tag = '', '', ''
                last_q = has_bull.group()
                _, end = has_bull.span()
                last_a = section[end:]
                last_tag = line_tag
            last_bottom = float(line_tag.rstrip('##').split('\t')[4])
            last_pn = line_pn
        if last_q:
            qai_list.append((last_q, last_a, *self.crop(last_tag, need_position=True)))
        return qai_list, tbls

    def get_tbls_info(self, tbls, tbl_index):
        if tbl_index >= len(tbls):
            return 1, 0, 0, 0, 0, '@@0\t0\t0\t0\t0##', ''
        tbl_pn = tbls[tbl_index][1][0][0]+1
        tbl_left = tbls[tbl_index][1][0][1]
        tbl_right = tbls[tbl_index][1][0][2]
        tbl_top = tbls[tbl_index][1][0][3]
        tbl_bottom = tbls[tbl_index][1][0][4]
        tbl_tag = "@@{}\t{:.1f}\t{:.1f}\t{:.1f}\t{:.1f}##" \
            .format(tbl_pn, tbl_left, tbl_right, tbl_top, tbl_bottom)
        _tbl_text = ''.join(tbls[tbl_index][0][1])
        return tbl_pn, tbl_left, tbl_right, tbl_top, tbl_bottom, tbl_tag, _tbl_text


class Docx(DocxParser):
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        img = paragraph._element.xpath('.//pic:pic')
        if not img:
            return None
        img = img[0]
        embed = img.xpath('.//a:blip/@r:embed')[0]
        related_part = document.part.related_parts[embed]
        image = related_part.image
        image = Image.open(BytesIO(image.blob)).convert('RGB')
        return image

    def __call__(self, filename, binary=None, from_page=0, to_page=100000, callback=None):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        last_answer, last_image = "", None
        question_stack, level_stack = [], []
        qai_list = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            question_level, p_text = 0, ''
            if from_page <= pn < to_page and p.text.strip():
                question_level, p_text = docx_question_level(p)
            if not question_level or question_level > 6: # not a question
                last_answer = f'{last_answer}\n{p_text}'
                current_image = self.get_picture(self.doc, p)
                last_image = concat_img(last_image, current_image)
            else:   # is a question
                if last_answer or last_image:
                    sum_question = '\n'.join(question_stack)
                    if sum_question:
                        qai_list.append((sum_question, last_answer, last_image))
                    last_answer, last_image = '', None

                i = question_level
                while question_stack and i <= level_stack[-1]:
                    question_stack.pop()
                    level_stack.pop()
                question_stack.append(p_text)
                level_stack.append(question_level)
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        if last_answer:
            sum_question = '\n'.join(question_stack)
            if sum_question:
                qai_list.append((sum_question, last_answer, last_image))

        tbls = []
        for tb in self.doc.tables:
            html= "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i+1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return qai_list, tbls


def rmPrefix(txt):
    return re.sub(
        r"^(问题|答案|回答|user|assistant|Q|A|Question|Answer|问|答)[\t:： ]+", "", txt.strip(), flags=re.IGNORECASE)


def beAdocPdf(d, q, a, eng, image, poss):
    qprefix = "Question: " if eng else "问题："
    aprefix = "Answer: " if eng else "回答："
    d["content_with_weight"] = "\t".join(
        [qprefix + rmPrefix(q), aprefix + rmPrefix(a)])
    d["content_ltks"] = rag_tokenizer.tokenize(q)
    d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
    d["image"] = image
    add_positions(d, poss)
    return d


def beAdocDocx(d, q, a, eng, image, row_num=-1):
    qprefix = "Question: " if eng else "问题："
    aprefix = "Answer: " if eng else "回答："
    d["content_with_weight"] = "\t".join(
        [qprefix + rmPrefix(q), aprefix + rmPrefix(a)])
    d["content_ltks"] = rag_tokenizer.tokenize(q)
    d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
    d["image"] = image
    if row_num >= 0:
        d["top_int"] = [row_num]
    return d


def beAdoc(d, q, a, eng, row_num=-1):
    qprefix = "Question: " if eng else "问题："
    aprefix = "Answer: " if eng else "回答："
    d["content_with_weight"] = "\t".join(
        [qprefix + rmPrefix(q), aprefix + rmPrefix(a)])
    d["content_ltks"] = rag_tokenizer.tokenize(q)
    d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
    if row_num >= 0:
        d["top_int"] = [row_num]
    return d


def mdQuestionLevel(s):
    match = re.match(r'#*', s)
    return (len(match.group(0)), s.lstrip('#').lstrip()) if match else (0, s)


def chunk(filename, binary=None, lang="Chinese", callback=None, **kwargs):
    """
        Excel and csv(txt) format files are supported.
        If the file is in excel format, there should be 2 column question and answer without header.
        And question column is ahead of answer column.
        And it's O.K if it has multiple sheets as long as the columns are rightly composed.

        If it's in csv format, it should be UTF-8 encoded. Use TAB as delimiter to separate question and answer.

        All the deformed lines will be ignored.
        Every pair of Q&A will be treated as a chunk.
    """
    eng = lang.lower() == "english"
    res = []
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    if re.search(r"\.xlsx?$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        excel_parser = Excel()
        for ii, (q, a) in enumerate(excel_parser(filename, binary, callback)):
            res.append(beAdoc(deepcopy(doc), q, a, eng, ii))
        return res

    elif re.search(r"\.(txt)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        lines = txt.split("\n")
        comma, tab = 0, 0
        for line in lines:
            if len(line.split(",")) == 2:
                comma += 1
            if len(line.split("\t")) == 2:
                tab += 1
        delimiter = "\t" if tab >= comma else ","

        fails = []
        question, answer = "", ""
        i = 0
        while i < len(lines):
            arr = lines[i].split(delimiter)
            if len(arr) != 2:
                if question:
                    answer += "\n" + lines[i]
                else:
                    fails.append(str(i+1))
            elif len(arr) == 2:
                if question and answer:
                    res.append(beAdoc(deepcopy(doc), question, answer, eng, i))
                question, answer = arr
            i += 1
            if len(res) % 999 == 0:
                callback(len(res) * 0.6 / len(lines), ("Extract Q&A: {}".format(len(res)) + (
                    f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))

        if question:
            res.append(beAdoc(deepcopy(doc), question, answer, eng, len(lines)))

        callback(0.6, ("Extract Q&A: {}".format(len(res)) + (
            f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))

        return res

    elif re.search(r"\.(csv)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        lines = txt.split("\n")
        delimiter = "\t" if any("\t" in line for line in lines) else ","

        fails = []
        question, answer = "", ""
        res = []
        reader = csv.reader(lines, delimiter=delimiter)

        for i, row in enumerate(reader):
            if len(row) != 2:
                if question:
                    answer += "\n" + lines[i]
                else:
                    fails.append(str(i + 1))
            elif len(row) == 2:
                if question and answer:
                    res.append(beAdoc(deepcopy(doc), question, answer, eng, i))
                question, answer = row
            if len(res) % 999 == 0:
                callback(len(res) * 0.6 / len(lines), ("Extract Q&A: {}".format(len(res)) + (
                    f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))

        if question:
            res.append(beAdoc(deepcopy(doc), question, answer, eng, len(list(reader))))

        callback(0.6, ("Extract Q&A: {}".format(len(res)) + (
            f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))
        return res

    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        pdf_parser = Pdf()
        qai_list, tbls = pdf_parser(filename if not binary else binary,
                                    from_page=0, to_page=10000, callback=callback)
        for q, a, image, poss in qai_list:
            res.append(beAdocPdf(deepcopy(doc), q, a, eng, image, poss))
        return res

    elif re.search(r"\.(md|markdown)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        lines = txt.split("\n")
        _last_question, last_answer = "", ""
        question_stack, level_stack = [], []
        code_block = False
        for index, line in enumerate(lines):
            if line.strip().startswith('```'):
                code_block = not code_block
            question_level, question = 0, ''
            if not code_block:
                question_level, question = mdQuestionLevel(line)

            if not question_level or question_level > 6: # not a question
                last_answer = f'{last_answer}\n{line}'
            else:   # is a question
                if last_answer.strip():
                    sum_question = '\n'.join(question_stack)
                    if sum_question:
                        res.append(beAdoc(deepcopy(doc), sum_question, markdown(last_answer, extensions=['markdown.extensions.tables']), eng, index))
                    last_answer = ''

                i = question_level
                while question_stack and i <= level_stack[-1]:
                    question_stack.pop()
                    level_stack.pop()
                question_stack.append(question)
                level_stack.append(question_level)
        if last_answer.strip():
            sum_question = '\n'.join(question_stack)
            if sum_question:
                res.append(beAdoc(deepcopy(doc), sum_question, markdown(last_answer, extensions=['markdown.extensions.tables']), eng, index))
        return res

    elif re.search(r"\.docx$", filename, re.IGNORECASE):
        docx_parser = Docx()
        qai_list, tbls = docx_parser(filename, binary,
                                    from_page=0, to_page=10000, callback=callback)
        res = tokenize_table(tbls, doc, eng)
        for i, (q, a, image) in enumerate(qai_list):
            res.append(beAdocDocx(deepcopy(doc), q, a, eng, image, i))
        return res

    raise NotImplementedError(
        "Excel, csv(txt), pdf, markdown and docx format files are supported.")


if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass
    chunk(sys.argv[1], from_page=0, to_page=10, callback=dummy)