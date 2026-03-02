#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

"""
Resume parsing module (aligned with SmartResume Pipeline architecture optimization)

Key optimizations (ref: arXiv:2510.09722):
    1. PDF text fusion: metadata + OCR dual-path extraction and fusion
    2. Layout-aware reconstruction: YOLOv10 layout segmentation + hierarchical sorting + line indexing
    3. Parallel task decomposition: basic info / work experience / education - 3-way parallel LLM extraction
    4. Index pointer mechanism: LLM returns line number ranges instead of generating full text, reducing hallucination
    5. Four-stage post-processing: source text re-extraction, domain normalization, context deduplication, source text validation

Compatibility:
    - chunk(filename, binary, callback, **kwargs) signature remains unchanged
    - Compatible with FACTORY[ParserType.RESUME.value] in task_executor.py
"""

import json
import re
import random
import datetime
import unicodedata
import concurrent.futures
from io import BytesIO
from typing import Optional

# tiktoken for long random string filtering (ref: SmartResume should_remove strategy)
try:
    import tiktoken
    _tiktoken_encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
except ImportError:
    _tiktoken_encoding = None

# Long random string pattern: 40+ char alphanumeric mixed strings (hash, token, tracking ID, etc.)
_LONG_RANDOM_PATTERN = re.compile(r'[a-zA-Z0-9\-~_]{40,}')

import logging as logger
from rag.nlp import rag_tokenizer
from deepdoc.parser.utils import get_text

# json_repair for fixing malformed JSON from LLM responses (ref: SmartResume fault-tolerance strategy)
try:
    import json_repair
except ImportError:
    json_repair = None

# YOLOv10 layout detector (lazy initialization to avoid loading model when unused)
_layout_recognizer = None


def _get_layout_recognizer():
    """
    Get YOLOv10 layout detector singleton (lazy loading)

    Uses the existing deepdoc LayoutRecognizer based on layout.onnx model.

    Returns:
        LayoutRecognizer instance, or None if loading fails
    """
    global _layout_recognizer
    if _layout_recognizer is None:
        try:
            from deepdoc.vision import LayoutRecognizer
            _layout_recognizer = LayoutRecognizer("layout")
            logger.info("YOLOv10 layout detector loaded successfully")
        except Exception as e:
            logger.warning(f"YOLOv10 layout detector loading failed, falling back to heuristic sorting: {e}")
            _layout_recognizer = False  # Mark as failed to avoid repeated attempts
    return _layout_recognizer if _layout_recognizer is not False else None

# ==================== Constants ====================

# Fields forbidden from being used as select fields in resume
FORBIDDEN_SELECT_FIELDS = [
    "name_pinyin_kwd", "edu_first_fea_kwd", "degree_kwd",
    "sch_rank_kwd", "edu_fea_kwd"
]

# Field name to description mapping (bilingual versions for chunk construction)
FIELD_MAP_ZH = {
    "name_kwd": "姓名/名字",
    "name_pinyin_kwd": "姓名拼音/名字拼音",
    "gender_kwd": "性别（男，女）",
    "age_int": "年龄/岁/年纪",
    "phone_kwd": "电话/手机/微信",
    "email_tks": "email/e-mail/邮箱",
    "position_name_tks": "职位/职能/岗位/职责",
    "expect_city_names_tks": "期望城市",
    "work_exp_flt": "工作年限/工作年份/N年经验/毕业了多少年",
    "corporation_name_tks": "最近就职(上班)的公司/上一家公司",
    "first_school_name_tks": "第一学历毕业学校",
    "first_degree_kwd": "第一学历",
    "highest_degree_kwd": "最高学历",
    "first_major_tks": "第一学历专业",
    "edu_first_fea_kwd": "第一学历标签",
    "degree_kwd": "过往学历",
    "major_tks": "学过的专业/过往专业",
    "school_name_tks": "学校/毕业院校",
    "sch_rank_kwd": "学校标签",
    "edu_fea_kwd": "教育标签",
    "corp_nm_tks": "就职过的公司/之前的公司/上过班的公司",
    "edu_end_int": "毕业年份",
    "industry_name_tks": "所在行业",
    "birth_dt": "生日/出生年份",
    "expect_position_name_tks": "期望职位/期望职能/期望岗位",
    "skill_tks": "技能/技术栈/编程语言/框架/工具",
    "language_tks": "语言能力/外语水平",
    "certificate_tks": "证书/资质/认证",
    "project_tks": "项目经验/项目名称",
    "work_desc_tks": "工作职责/工作描述",
    "project_desc_tks": "项目描述/项目职责",
    "self_evaluation_tks": "自我评价/个人优势/个人总结",
}

FIELD_MAP_EN = {
    "name_kwd": "Name",
    "name_pinyin_kwd": "Name Pinyin",
    "gender_kwd": "Gender (Male, Female)",
    "age_int": "Age",
    "phone_kwd": "Phone/Mobile/WeChat",
    "email_tks": "Email",
    "position_name_tks": "Position/Title/Role",
    "expect_city_names_tks": "Preferred City",
    "work_exp_flt": "Years of Experience",
    "corporation_name_tks": "Most Recent Company",
    "first_school_name_tks": "First Degree School",
    "first_degree_kwd": "First Degree",
    "highest_degree_kwd": "Highest Degree",
    "first_major_tks": "First Degree Major",
    "edu_first_fea_kwd": "First Degree Tag",
    "degree_kwd": "Past Degrees",
    "major_tks": "Past Majors",
    "school_name_tks": "School/University",
    "sch_rank_kwd": "School Tag",
    "edu_fea_kwd": "Education Tag",
    "corp_nm_tks": "Past Companies",
    "edu_end_int": "Graduation Year",
    "industry_name_tks": "Industry",
    "birth_dt": "Date of Birth",
    "expect_position_name_tks": "Preferred Position/Role",
    "skill_tks": "Skills/Tech Stack/Languages/Frameworks/Tools",
    "language_tks": "Language Proficiency",
    "certificate_tks": "Certificates/Qualifications",
    "project_tks": "Project Experience/Project Name",
    "work_desc_tks": "Job Responsibilities/Description",
    "project_desc_tks": "Project Description/Responsibilities",
    "self_evaluation_tks": "Self-Evaluation/Personal Strengths/Summary",
}


def _is_english(lang: str) -> bool:
    """Determine if the language parameter indicates English"""
    return lang.lower() in ("english", "en")


def get_field_map(lang: str) -> dict:
    """Get the corresponding field mapping based on language parameter"""
    return FIELD_MAP_EN if _is_english(lang) else FIELD_MAP_ZH


# Backward compatible: default to Chinese version
FIELD_MAP = FIELD_MAP_ZH


# ==================== Parallel LLM Extraction Prompt Templates ====================
# Ref: SmartResume task decomposition strategy, splitting extraction into independent subtasks
# Each prompt ends with /no_think marker to suppress reasoning model's thinking output
# Prompts loaded from md files under rag/prompts/, supporting bilingual versions

from rag.prompts.template import load_prompt


def _load_resume_prompt(name: str, lang: str) -> str:
    """Load the corresponding version of resume prompt template based on language parameter

    Args:
        name: Prompt name (without language suffix), e.g. "resume_system"
        lang: Language parameter, e.g. "Chinese" or "English"
    Returns:
        Prompt template string
    """
    suffix = "_en" if _is_english(lang) else ""
    return load_prompt(f"{name}{suffix}")


def get_system_prompt(lang: str) -> str:
    """Get system prompt"""
    return _load_resume_prompt("resume_system", lang)


def get_basic_info_prompt(lang: str) -> str:
    """Get basic info extraction prompt"""
    return _load_resume_prompt("resume_basic_info", lang)


def get_work_exp_prompt(lang: str) -> str:
    """Get work experience extraction prompt"""
    return _load_resume_prompt("resume_work_exp", lang)


def get_education_prompt(lang: str) -> str:
    """Get education background extraction prompt"""
    return _load_resume_prompt("resume_education", lang)


def get_project_exp_prompt(lang: str) -> str:
    """Get project experience extraction prompt"""
    return _load_resume_prompt("resume_project_exp", lang)


# Backward compatible: default Chinese version constants (for possible external direct references)
SYSTEM_PROMPT = load_prompt("resume_system")
BASIC_INFO_PROMPT = load_prompt("resume_basic_info")
WORK_EXP_PROMPT = load_prompt("resume_work_exp")
EDUCATION_PROMPT = load_prompt("resume_education")
PROJECT_EXP_PROMPT = load_prompt("resume_project_exp")

# LLM call max retry count (ref: SmartResume retry strategy)
_LLM_MAX_RETRIES = 2


def _normalize_whitespace(text: str) -> str:
    """
    Unicode whitespace normalization (ref: SmartResume _clean_text_content)

    Replaces various Unicode spaces (\u00A0 non-breaking space, \u3000 fullwidth space,
    \u2000-\u200A various width spaces, etc.) with regular spaces,
    then applies NFKC normalization (fullwidth to halfwidth) and merges consecutive spaces.

    Args:
        text: Original text
    Returns:
        Normalized text
    """
    if not text:
        return ""
    # NFKC normalization (fullwidth to halfwidth, etc.)
    text = unicodedata.normalize('NFKC', text)
    # Unify various Unicode spaces to regular space
    text = re.sub(
        r'[\u0020\u00A0\u1680\u2000-\u200A\u2028\u2029\u202F\u205F\u3000\u00A7]',
        ' ', text
    )
    # Merge consecutive spaces
    text = re.sub(r' {2,}', ' ', text)
    return text.strip()


def _should_remove_random_str(match: re.Match) -> bool:
    """
    Determine if a matched long string is a meaningless random string (ref: SmartResume should_remove)

    Uses tiktoken encoding to judge: if token count exceeds 50% of original char count,
    it indicates a meaningless random string (hash, token, tracking ID, etc.) that should be removed.
    Normal English words have high token encoding efficiency, with token count far less than char count.

    Args:
        match: Regex match object
    Returns:
        True means it should be removed
    """
    if _tiktoken_encoding is None:
        # When tiktoken is unavailable, use simple heuristic: case/digit alternation frequency
        s = match.group(0)
        changes = sum(
            1 for i in range(1, len(s))
            if s[i].isdigit() != s[i-1].isdigit()
            or (s[i].isalpha() and s[i-1].isalpha() and s[i].isupper() != s[i-1].isupper())
        )
        return changes / len(s) > 0.3
    encoded = _tiktoken_encoding.encode(match.group(0))
    return len(encoded) > len(match.group(0)) * 0.5


def _clean_line_content(text: str) -> str:
    """
    Clean single line text content (Unicode normalization + long random string filtering)

    Args:
        text: Original line text
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    # Unicode whitespace normalization
    text = _normalize_whitespace(text)
    # Filter long random strings (hash, token and other meaningless content)
    text = _LONG_RANDOM_PATTERN.sub(
        lambda m: '' if _should_remove_random_str(m) else m.group(0),
        text
    )
    # Clean up extra spaces after filtering
    text = re.sub(r' {2,}', ' ', text).strip()
    return text


# ==================== Phase 1: PDF Text Fusion and Layout Reconstruction ====================




def _is_noise_char(obj: dict) -> bool:
    """
    Determine if a PDF character object is a decorative layer noise character

    Uses a "body text whitelist" strategy instead of enumerating noise features,
    to handle noise patterns from different resume templates:

    Two reliable features of body text characters (either one means body text):
    1. Embedded font: Font name format is XXXXXX+FontName (contains '+'),
       indicating the font is embedded in the PDF, chosen by the document author
    2. Structure tag: Has PDF Tagged Structure tags (e.g., Span, P, NonStruct, etc.),
       indicating the character belongs to the document's semantic structure tree

    Common features of noise characters:
    - Uses system fonts (e.g., Helvetica, Arial), font name doesn't contain '+'
    - No structure tags (tag is None or non-semantic tags like 'OC')
    - Common in resume template background decorations, watermarks, tracking marks

    Args:
        obj: pdfplumber character/text object dictionary
    Returns:
        True means it's a noise character that should be filtered
    """
    # Whitelist condition 1: Embedded font (font name contains '+' prefix)
    fontname = obj.get("fontname", "")
    if "+" in fontname:
        return False  # Embedded font = body content

    # Whitelist condition 2: Has PDF structure tag
    tag = obj.get("tag")
    if tag in ("Span", "NonStruct", "P", "H1", "H2", "H3", "H4", "H5", "H6",
               "TD", "TH", "LI", "L", "Table", "TR", "Figure", "Caption"):
        return False  # Has semantic structure tag = body content

    # Doesn't meet any whitelist condition, treat as noise
    return True



def _extract_metadata_text(binary: bytes) -> list[dict]:
    """
    Extract text blocks from PDF metadata (with coordinate info)

    Strategy:
    1. Use whitelist strategy to filter decorative layer noise chars (embedded font or structure tag = body text)
    2. Safe fallback: if filtered chars are less than 30% of original, skip filtering to avoid false positives
    3. Use extract_words for word-level extraction (with real coordinates)
    4. Aggregate adjacent words into line-level text blocks by Y coordinate
    5. Additionally extract table content (many resumes use table layouts)

    Args:
        binary: PDF file binary content
    Returns:
        List of text blocks, each containing text, x0, top, x1, bottom, page fields
    """
    try:
        import pdfplumber
        blocks = []
        with pdfplumber.open(BytesIO(binary)) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                page_width = page.width or 600

                # Filter decorative layer noise chars (whitelist strategy based on embedded font + structure tag)
                # Safe fallback: if filtered chars are less than 30% of original, the PDF's body text
                # may use non-embedded fonts without structure tags, skip filtering to avoid false positives
                try:
                    original_char_count = len(page.chars)
                    filtered_page = page.filter(
                        lambda obj: not _is_noise_char(obj)
                    )
                    filtered_char_count = len(filtered_page.chars)
                    if original_char_count > 0 and filtered_char_count < original_char_count * 0.3:
                        # Filtered out over 70% of chars, likely false positives, fall back to original page
                        filtered_page = page
                except Exception:
                    filtered_page = page

                # Use extract_words for extraction (with real coordinates)
                words = []
                try:
                    words = filtered_page.extract_words(
                        keep_blank_chars=False, use_text_flow=True
                    )
                except Exception:
                    pass

                if words:
                    # Aggregate adjacent words into line-level text blocks by Y coordinate
                    # Words on the same line: top coordinate difference within threshold
                    line_threshold = 5  # Y coordinate difference threshold (unit: PDF points)
                    current_line_words = [words[0]]

                    def _flush_line(line_words):
                        """Merge words in a line into a single text block"""
                        # Sort by x0 to ensure left-to-right order
                        line_words.sort(key=lambda w: float(w.get("x0", 0)))
                        texts = []
                        for w in line_words:
                            texts.append(w.get("text", ""))
                        merged_text = " ".join(texts)
                        if not merged_text.strip():
                            return None
                        return {
                            "text": merged_text.strip(),
                            "x0": float(min(w.get("x0", 0) for w in line_words)),
                            "top": float(min(w.get("top", 0) for w in line_words)),
                            "x1": float(max(w.get("x1", 0) for w in line_words)),
                            "bottom": float(max(w.get("bottom", 0) for w in line_words)),
                            "page": page_idx,
                        }

                    for w in words[1:]:
                        w_top = float(w.get("top", 0))
                        cur_top = float(current_line_words[0].get("top", 0))
                        if abs(w_top - cur_top) <= line_threshold:
                            current_line_words.append(w)
                        else:
                            block = _flush_line(current_line_words)
                            if block:
                                blocks.append(block)
                            current_line_words = [w]

                    # Process the last line
                    if current_line_words:
                        block = _flush_line(current_line_words)
                        if block:
                            blocks.append(block)
                else:
                    # Fall back to extract_text when extract_words fails
                    page_text = None
                    try:
                        page_text = page.extract_text()
                    except Exception:
                        pass
                    if page_text and page_text.strip():
                        raw_lines = page_text.split("\n")
                        line_height = 16
                        for i, line in enumerate(raw_lines):
                            cleaned = line.strip()
                            if not cleaned:
                                continue
                            blocks.append({
                                "text": cleaned,
                                "x0": 0,
                                "top": i * line_height,
                                "x1": page_width,
                                "bottom": i * line_height + line_height - 2,
                                "page": page_idx,
                            })

                # Extract table content from the page
                # Many resumes use table layouts (e.g., personal info section), extract_words may miss table structure
                try:
                    tables = page.extract_tables()
                    if tables:
                        page_blocks = [b for b in blocks if b["page"] == page_idx]
                        max_top = max((b["top"] for b in page_blocks), default=0) + 20
                        row_height = 16

                        for table in tables:
                            for row in table:
                                if not row:
                                    continue
                                cells = [str(c).strip() for c in row if c and str(c).strip()]
                                if not cells:
                                    continue
                                row_text = " | ".join(cells)
                                # Dedup: check if table content was already extracted by extract_words
                                is_dup = False
                                for pb in page_blocks:
                                    if all(c in pb["text"] for c in cells[:2]):
                                        is_dup = True
                                        break
                                if is_dup:
                                    continue
                                blocks.append({
                                    "text": row_text,
                                    "x0": 0,
                                    "top": max_top,
                                    "x1": page_width,
                                    "bottom": max_top + row_height - 2,
                                    "page": page_idx,
                                })
                                max_top += row_height
                except Exception as e:
                    logger.debug(f"PDF table extraction skipped (page {page_idx}): {e}")
        return blocks
    except Exception as e:
        logger.warning(f"PDF metadata extraction failed: {e}")
        return []

def _extract_ocr_text(binary: bytes, meta_blocks: list[dict] | None = None) -> list[dict]:
    """
    Extract OCR text blocks using blackout strategy (with coordinate info).

    Strategy (ref: SmartResume):
    1. Render PDF pages to images
    2. Black out regions already extracted by metadata
    3. Run OCR on the blacked-out image, only recognizing content metadata missed
    4. Eliminates duplication at source, no IoU dedup needed downstream

    Args:
        binary: PDF file binary content
        meta_blocks: Text blocks from metadata extraction, used to black out existing text regions
    Returns:
        List of text blocks, each containing text, x0, top, x1, bottom, page fields
    """
    if meta_blocks is None:
        meta_blocks = []
    try:
        import pdfplumber
        from deepdoc.vision.ocr import OCR
        import numpy as np

        ocr = OCR()
        blocks = []

        with pdfplumber.open(BytesIO(binary)) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                # Render page to image (resolution=216 = 3x scale, since PDF default is 72 DPI)
                img = page.to_image(resolution=216)
                page_img = np.array(img.annotated)

                # Scale factor from PDF coordinates to image coordinates
                pdf_to_img_scale = 216.0 / 72.0  # = 3.0

                # Black out metadata-extracted text regions before OCR
                page_meta_blocks = [b for b in meta_blocks if b.get("page") == page_idx]
                if page_meta_blocks:
                    page_img = _blackout_text_regions(page_img, meta_blocks, page_idx, pdf_to_img_scale)

                ocr_result = ocr(page_img)
                if not ocr_result:
                    continue
                for box_info in ocr_result:
                    if isinstance(box_info, (list, tuple)) and len(box_info) >= 2:
                        coords = box_info[0]  # Coordinate points
                        text_info = box_info[1]
                        text = text_info[0] if isinstance(text_info, (list, tuple)) else str(text_info)
                        if text.strip() and isinstance(coords, (list, tuple)) and len(coords) >= 4:
                            # Extract bounding box from four corner points
                            xs = [p[0] for p in coords if isinstance(p, (list, tuple))]
                            ys = [p[1] for p in coords if isinstance(p, (list, tuple))]
                            if xs and ys:
                                blocks.append({
                                    "text": text.strip(),
                                    "x0": min(xs), "top": min(ys),
                                    "x1": max(xs), "bottom": max(ys),
                                    "page": page_idx,
                                })
        return blocks
    except Exception as e:
        logger.warning(f"OCR extraction failed: {e}")
        return []


def _fuse_text_blocks(meta_blocks: list[dict], ocr_blocks: list[dict]) -> list[dict]:
    """
    Fuse PDF metadata text and OCR text (blackout strategy version).

    Since the OCR phase already blacks out metadata-extracted regions, OCR only recognizes
    content that metadata missed. Therefore this function only needs to:
    1. Filter out garbled blocks from metadata
    2. Directly merge valid metadata blocks and OCR blocks (no IoU dedup needed)

    Args:
        meta_blocks: Text blocks from metadata extraction
        ocr_blocks: Text blocks from OCR extraction (already deduplicated via blackout strategy)
    Returns:
        Fused text block list
    """
    if not ocr_blocks:
        return meta_blocks
    if not meta_blocks:
        return ocr_blocks

    # Filter out garbled blocks from metadata
    valid_meta = []
    garbled_count = 0
    for b in meta_blocks:
        if _is_valid_line(b.get("text", "")):
            valid_meta.append(b)
        else:
            garbled_count += 1

    if garbled_count:
        logger.info(f"Detected {garbled_count} garbled blocks in metadata, filtered out")

    # Under blackout strategy, OCR won't re-recognize existing text, just merge directly
    fused = valid_meta + ocr_blocks
    return fused




def _layout_aware_reorder(blocks: list[dict]) -> list[dict]:
    """
    Layout-aware hierarchical sorting (ref: SmartResume Hierarchical Re-ordering)

    Two-level sorting strategy:
    1. Inter-segment sorting: first by page number, then by Y coordinate (top to bottom), same row by X coordinate (left to right)
    2. Intra-segment sorting: within each logical segment, sort by reading order

    For multi-column resumes, detect column positions by clustering X coordinates,
    then sort by column order.

    Args:
        blocks: Text block list (with coordinate info)
    Returns:
        Sorted text block list
    """
    if not blocks:
        return blocks

    # Group by page
    pages = {}
    for b in blocks:
        pg = b.get("page", 0)
        pages.setdefault(pg, []).append(b)

    sorted_blocks = []
    for pg in sorted(pages.keys()):
        page_blocks = pages[pg]

        # Detect multi-column layout: by X coordinate median
        if len(page_blocks) > 5:
            x_centers = [(b["x0"] + b["x1"]) / 2 for b in page_blocks]
            x_min, x_max = min(x_centers), max(x_centers)
            page_width = x_max - x_min if x_max > x_min else 1

            # Simple two-column detection: if text blocks are clearly distributed on left and right sides
            mid_x = (x_min + x_max) / 2
            left_count = sum(1 for x in x_centers if x < mid_x - page_width * 0.1)
            right_count = sum(1 for x in x_centers if x > mid_x + page_width * 0.1)

            if left_count > 3 and right_count > 3:
                # Multi-column layout: left column first then right column, each column top to bottom
                left_blocks = [b for b in page_blocks if (b["x0"] + b["x1"]) / 2 < mid_x]
                right_blocks = [b for b in page_blocks if (b["x0"] + b["x1"]) / 2 >= mid_x]
                left_blocks.sort(key=lambda b: (b["top"], b["x0"]))
                right_blocks.sort(key=lambda b: (b["top"], b["x0"]))
                sorted_blocks.extend(left_blocks)
                sorted_blocks.extend(right_blocks)
                continue

        # Single-column layout: top to bottom, same row left to right
        page_blocks.sort(key=lambda b: (b["top"], b["x0"]))
        sorted_blocks.extend(page_blocks)

    return sorted_blocks


def _layout_detect_reorder(blocks: list[dict], binary: bytes) -> list[dict]:
    """
    Use YOLOv10 layout detection for layout-aware sorting (ref: SmartResume Layout-Aware Reordering)

    Flow:
    1. Render each PDF page as an image
    2. Use YOLOv10 to detect layout regions (title, body, table, etc.)
    3. Assign text blocks to detected layout regions
    4. Sort hierarchically by region position (region center Y -> region center X -> block Y -> block X)

    Automatically falls back to heuristic sorting on detection failure.

    Args:
        blocks: Text block list (with coordinate info)
        binary: PDF file binary content (for rendering page images)
    Returns:
        Sorted text block list
    """
    if not blocks:
        return blocks

    recognizer = _get_layout_recognizer()
    if recognizer is None:
        logger.info("Layout detector unavailable, falling back to heuristic sorting")
        return _layout_aware_reorder(blocks)

    try:
        import pdfplumber
        # Group text blocks by page
        pages_blocks = {}
        for b in blocks:
            pg = b.get("page", 0)
            pages_blocks.setdefault(pg, []).append(b)

        # Render each page as image and prepare OCR-format input
        page_indices = sorted(pages_blocks.keys())
        image_list = []
        ocr_res_per_page = []

        with pdfplumber.open(BytesIO(binary)) as pdf:
            for pg in page_indices:
                if pg >= len(pdf.pages):
                    continue
                page = pdf.pages[pg]
                # Render as PIL Image (scale_factor=3 matches LayoutRecognizer default)
                pil_img = page.to_image(resolution=72 * 3).annotated
                image_list.append(pil_img)

                # Convert page text blocks to LayoutRecognizer required format
                page_bxs = []
                for b in pages_blocks[pg]:
                    page_bxs.append({
                        "x0": float(b["x0"]),
                        "top": float(b["top"]),
                        "x1": float(b["x1"]),
                        "bottom": float(b["bottom"]),
                        "text": b["text"],
                        "page": pg,
                    })
                ocr_res_per_page.append(page_bxs)

        if not image_list:
            return _layout_aware_reorder(blocks)

        # Call YOLOv10 layout detection + text block annotation
        # LayoutRecognizer.__call__ tags each text block with layout_type and layoutno
        tagged_blocks, page_layouts = recognizer(
            image_list, ocr_res_per_page, scale_factor=3, thr=0.2, drop=False
        )

        if not tagged_blocks:
            logger.warning("Layout detection returned no results, falling back to heuristic sorting")
            return _layout_aware_reorder(blocks)

        # Sort by layoutno groups:
        # 1. Blocks with layoutno sorted by region position (region Y -> region X -> block Y -> block X)
        # 2. Blocks without layoutno sorted by original coordinates
        def _sort_key(b):
            layout_type = b.get("layout_type", "")
            # header sorts first, footer sorts last
            if layout_type == "header":
                return (0, 0, b.get("top", 0), b.get("x0", 0))
            if layout_type == "footer":
                return (9999, 0, b.get("top", 0), b.get("x0", 0))
            # Others sorted by top -> x0
            return (1, b.get("top", 0), b.get("x0", 0), 0)

        tagged_blocks.sort(key=_sort_key)

        # Restore page field (LayoutRecognizer may not preserve it)
        for b in tagged_blocks:
            if "page" not in b:
                b["page"] = 0

        logger.info(f"YOLOv10 layout detection complete, {len(tagged_blocks)} text blocks total, "
                    f"detected {sum(len(pl) for pl in page_layouts)} layout regions")
        return tagged_blocks

    except Exception as e:
        logger.warning(f"YOLOv10 layout detection sorting failed, falling back to heuristic sorting: {e}")
        return _layout_aware_reorder(blocks)



def _build_indexed_text(blocks: list[dict]) -> tuple[str, list[str], list[dict]]:
    """
    Build indexed text with line numbers (ref: SmartResume Indexed Linearization)

    Merges sorted text blocks into lines and adds a unique index number to each line.
    Includes garbled line filtering logic and field label split repair.
    Also preserves coordinate info for each line, used for writing position_int etc. to chunks.

    Args:
        blocks: Sorted text block list
    Returns:
        (indexed_text, lines, line_positions) tuple:
        - indexed_text: Text string with line numbers
        - lines: Original line text list (without line numbers)
        - line_positions: Coordinate info for each line, format:
          {"page": int, "x0": float, "x1": float, "top": float, "bottom": float}
    """
    if not blocks:
        return "", [], []

    # Merge adjacent text blocks into lines (based on Y coordinate proximity)
    # Also record bounding box for each line (outer bounding rectangle of all blocks)
    raw_lines = []
    raw_positions = []  # Coordinates for each line
    current_line_parts = []
    current_line_blocks = []  # All blocks in current line
    current_top = blocks[0].get("top", 0)
    threshold = 10  # Y coordinate difference threshold

    def _merge_line_position(line_blocks: list[dict]) -> dict:
        """Merge coordinates of all blocks in a line into outer bounding rectangle"""
        return {
            "page": line_blocks[0].get("page", 0),
            "x0": min(b.get("x0", 0) for b in line_blocks),
            "x1": max(b.get("x1", 0) for b in line_blocks),
            "top": min(b.get("top", 0) for b in line_blocks),
            "bottom": max(b.get("bottom", 0) for b in line_blocks),
        }

    for b in blocks:
        if abs(b.get("top", 0) - current_top) > threshold and current_line_parts:
            raw_lines.append(" ".join(current_line_parts))
            raw_positions.append(_merge_line_position(current_line_blocks))
            current_line_parts = []
            current_line_blocks = []
            current_top = b.get("top", 0)
        current_line_parts.append(b["text"])
        current_line_blocks.append(b)

    if current_line_parts:
        raw_lines.append(" ".join(current_line_parts))
        raw_positions.append(_merge_line_position(current_line_blocks))

    # Filter empty and garbled lines (sync filter coordinates)
    lines = []
    line_positions = []
    for line, pos in zip(raw_lines, raw_positions):
        # Unicode normalization + long random string filtering (ref: SmartResume _clean_text_content)
        line = _clean_line_content(line)
        if not line:
            continue
        # Garbled detection: skip if valid chars (Chinese/ASCII letters/digits/common punctuation) ratio is too low
        if not _is_valid_line(line):
            continue
        lines.append(line)
        line_positions.append(pos)

    # Fix field label split issues
    # Coordinates are not affected, keep original positions
    lines = _fix_split_labels(lines)

    # Build indexed text with line numbers
    indexed_parts = [f"[{i}]: {line}" for i, line in enumerate(lines)]
    indexed_text = "\n".join(indexed_parts)

    return indexed_text, lines, line_positions

def _is_valid_line(line: str) -> bool:
    """
    Check if a text line is valid content (not garbled)

    Multi-dimensional detection:
    1. Valid character ratio (Chinese, ASCII alphanumeric, common punctuation)
    2. Single-character spacing anomaly detection (PDF custom font mapping causing "O U W Z_W V 2" pattern)
    3. Consecutive meaningless alphanumeric sequence detection

    Args:
        line: Text line to check
    Returns:
        True means valid line, False means garbled line
    """
    if len(line) <= 3:
        # Short lines may be valid content like names, keep them
        return True

    cid_count = len(re.findall(r'\(cid:\d+\)', line))
    if cid_count >= 3:
        return False
    # Valid characters: Chinese (incl. extension), ASCII alphanumeric, common punctuation and spaces, fullwidth chars, CJK punctuation
    valid_chars = re.findall(
        r'[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff'
        r'a-zA-Z0-9\s@.,:;!?()（）【】\-_/\\|·•'
        r'、，。：；！？\u201c\u201d\u2018\u2019《》'
        r'\uff01-\uff5e'
        r'\u3000-\u303f'
        r'#%&+=~`\u00b7\u2022\u2013\u2014'
        r']',
        line
    )
    ratio = len(valid_chars) / len(line) if len(line) > 0 else 0
    if ratio < 0.5:
        return False

    # Detect PDF custom font mapping causing single-character spacing anomaly pattern
    # Feature: lots of "single letter space single letter space" sequences, e.g. "O U W Z_W V 2 X 3"
    # Stats: ratio of space-separated single chars among non-space chars
    spaced_singles = re.findall(r'(?:^|\s)([a-zA-Z0-9])(?:\s|$)', line)
    non_space_len = len(line.replace(" ", ""))
    if non_space_len > 5 and len(spaced_singles) > 0:
        # If ratio of space-separated single chars to non-space chars is too high, classify as garbled
        single_ratio = len(spaced_singles) / non_space_len
        if single_ratio > 0.3:
            return False

    # Detect consecutive meaningless mixed-case alphanumeric sequences (e.g. "UJqZX9V2")
    # Normal English words don't have such frequent case alternation patterns
    garbled_seqs = re.findall(r'[a-zA-Z0-9]{4,}', line.replace(" ", ""))
    if garbled_seqs:
        garbled_count = 0
        for seq in garbled_seqs:
            # Count case alternations
            case_changes = sum(
                1 for i in range(1, len(seq))
                if (seq[i].isupper() != seq[i-1].isupper() and seq[i].isalpha() and seq[i-1].isalpha())
                or (seq[i].isdigit() != seq[i-1].isdigit())
            )
            # Too high alternation frequency = garbled sequence (normal words like "Spring" have only 1 alternation)
            if len(seq) >= 4 and case_changes / len(seq) > 0.5:
                garbled_count += 1
        # If garbled sequence ratio is too high
        if len(garbled_seqs) > 0 and garbled_count / len(garbled_seqs) > 0.4:
            return False

    return True


def _fix_split_labels(lines: list[str]) -> list[str]:
    """
    Fix field label split issues

    Some PDF layouts split field labels across line start/end, e.g.:
    - "名：陈晓俐 姓" -> should be fixed to "姓名：陈晓俐"
    - "别：男 性" -> should be fixed to "性别：男"

    Args:
        lines: Original line text list
    Returns:
        Fixed line text list
    """
    # Common split field label patterns: (line-end part, line-start part) -> full label
    split_patterns = {
        ("姓", "名"): "姓名",
        ("性", "别"): "性别",
        ("年", "龄"): "年龄",
        ("电", "话"): "电话",
        ("邮", "箱"): "邮箱",
        ("学", "历"): "学历",
        ("专", "业"): "专业",
        ("地", "址"): "地址",
        ("籍", "贯"): "籍贯",
        ("民", "族"): "民族",
    }

    fixed = []
    for line in lines:
        # Detect in-line split patterns: "X：content Y" where (Y, X) is a split pair
        for (suffix_char, prefix_char), full_label in split_patterns.items():
            # Pattern: "prefix_char：content suffix_char" (first half at line start, second half at line end)
            pattern = rf'^({re.escape(prefix_char)})\s*[:：]\s*(.+?)\s+{re.escape(suffix_char)}\s*$'
            m = re.match(pattern, line)
            if m:
                content = m.group(2).strip()
                line = f"{full_label}：{content}"
                break
            # Pattern: "suffix_char content prefix_char：" (second half at line start, first half at line end)
            pattern2 = rf'^{re.escape(suffix_char)}\s*[:：]?\s*(.+?)\s+{re.escape(prefix_char)}\s*$'
            m2 = re.match(pattern2, line)
            if m2:
                content = m2.group(1).strip()
                line = f"{full_label}：{content}"
                break
        fixed.append(line)
    return fixed





def extract_text(filename: str, binary: bytes) -> tuple[str, list[str], list[dict]]:
    """
    Extract text content based on file type (Pipeline Phase 1).

    PDF files use dual-path fusion + layout reconstruction + line indexing.
    Other formats fall back to simple text extraction.

    Args:
        filename: File name
        binary: File binary content
    Returns:
        (indexed_text, lines, line_positions) tuple:
        - indexed_text: Text with line number indices
        - lines: List of original line texts
        - line_positions: List of per-line coordinate info (empty list for non-PDF formats)
    """
    fname_lower = filename.lower()

    try:
        if fname_lower.endswith(".pdf"):
            # Dual-path extraction
            meta_blocks = _extract_metadata_text(binary)
            ocr_blocks = []

            # Determine whether OCR supplementation is needed:
            # 1. Metadata text too short (< 100 chars)
            # 2. High garbled text ratio in metadata (caused by custom font mapping)
            meta_text_len = sum(len(b["text"]) for b in meta_blocks)
            need_ocr = False

            if meta_text_len < 100:
                logger.info("PDF metadata text too short, enabling OCR supplementation")
                need_ocr = True
            else:
                # Check metadata text quality: calculate valid line ratio
                # If many lines are judged as garbled by _is_valid_line, the PDF font mapping has issues
                valid_line_count = 0
                total_line_count = 0
                for b in meta_blocks:
                    text = b.get("text", "").strip()
                    if not text:
                        continue
                    total_line_count += 1
                    if _is_valid_line(text):
                        valid_line_count += 1
                if total_line_count > 0:
                    valid_ratio = valid_line_count / total_line_count
                    if valid_ratio < 0.6:
                        logger.info(
                            f"PDF metadata text quality low (valid line ratio {valid_ratio:.1%}), enabling OCR supplementation"
                        )
                        need_ocr = True

            if need_ocr:
                # Blackout strategy: black out metadata-extracted regions before OCR
                ocr_blocks = _extract_ocr_text(binary, meta_blocks=meta_blocks)

            # Text fusion
            fused_blocks = _fuse_text_blocks(meta_blocks, ocr_blocks)

            # Layout-aware sorting (prefer YOLOv10 layout detection, fall back to heuristic on failure)
            sorted_blocks = _layout_detect_reorder(fused_blocks, binary)

            # Build line-indexed text (with coordinate info)
            return _build_indexed_text(sorted_blocks)

        elif fname_lower.endswith(".docx"):
            from docx import Document
            doc = Document(BytesIO(binary))
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            # Extract table content from DOCX
            # Reference: table handling in naive.py Docx class
            # Many resumes use table layouts for personal info; iterating only paragraphs would miss this content
            for table in doc.tables:
                for row in table.rows:
                    cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            cells.append(cell_text)
                    if not cells:
                        continue
                    row_text = " | ".join(cells)
                    # Deduplicate: skip if this row text already exists in lines
                    if row_text not in lines:
                        lines.append(row_text)

            indexed = "\n".join(f"[{i}]: {line}" for i, line in enumerate(lines))
            # DOCX has no coordinate info, return empty list
            return indexed, lines, []

        else:
            text = get_text(filename, binary)
            lines = [line.strip() for line in text.split("\n") if line.strip()]
            indexed = "\n".join(f"[{i}]: {line}" for i, line in enumerate(lines))
            return indexed, lines, []

    except Exception:
        logger.exception(f"Text extraction failed: {filename}")
        return "", [], []


# ==================== Phase 2: Parallel LLM Structured Extraction ====================


def _clean_llm_json_response(response: str) -> str:
    """
    Clean LLM JSON response.

    Uses SmartResume's lightweight string extraction strategy:
    1. Remove markdown code block markers
    2. Remove <think>...</think> thinking tags (reasoning models may output these)
    3. text.find("{") and text.rfind("}") to locate valid JSON block

    Args:
        response: Raw LLM response text
    Returns:
        Cleaned JSON string
    """
    text = response.strip()
    # Remove markdown code block markers
    text = text.replace("```json", "").replace("```", "").strip()
    # Remove reasoning model thinking tags
    text = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL).strip()
    # Clean escaped quotes (SmartResume's approach)
    text = text.replace('\\"', '"')
    # SmartResume strategy: locate first { and last }
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        return text[start:end + 1]
    return text


def _parse_json_with_repair(text: str) -> dict:
    """
    Parse JSON string, attempt repair on failure (ref SmartResume's json_repair strategy).

    Repair strategies:
    1. Standard json.loads
    2. Replace Python-style booleans/None
    3. Use json_repair library

    Args:
        text: JSON string
    Returns:
        Parsed dictionary
    Raises:
        json.JSONDecodeError: Raised when all repair strategies fail
    """
    # First attempt: standard parsing
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Second attempt: replace Python-style values (ref SmartResume)
    repaired = text.replace("'", '"')
    repaired = repaired.replace('True', 'true')
    repaired = repaired.replace('False', 'false')
    repaired = repaired.replace('None', 'null')
    try:
        return json.loads(repaired)
    except json.JSONDecodeError:
        pass

    # Third attempt: use json_repair library
    if json_repair is not None:
        try:
            return json_repair.loads(text)
        except Exception:
            pass

    # All strategies failed
    raise json.JSONDecodeError("All JSON repair strategies failed", text, 0)


def _call_llm(prompt: str, tenant_id , lang: str) -> Optional[dict]:
    """
    Call LLM and parse JSON response (ref SmartResume's retry + fault-tolerance strategy).

    Retry mechanism:
    - Retry up to _LLM_MAX_RETRIES times
    - On retry, increase temperature and randomize seed for output diversity
    - Use json_repair on JSON parse failure

    Args:
        prompt: User prompt
        lang: Language
    Returns:
        Parsed dictionary, or None on failure

    """
    try:
        from api.db.services.llm_service import LLMBundle
        from common.constants import LLMType

        llm =  LLMBundle(tenant_id, LLMType.CHAT, lang=lang)

        for attempt in range(_LLM_MAX_RETRIES + 1):
            try:
                # Increase temperature on retry for diversity (ref SmartResume)
                temperature = 0.1 if attempt == 0 else 1.0
                gen_conf = {"temperature": temperature, "max_tokens": 2048}
                if attempt > 0:
                    gen_conf["seed"] = random.randint(0, 1000000)

                response = llm.chat(
                    system=get_system_prompt(lang),
                    history=[{"role": "user", "content": prompt}],
                    gen_conf=gen_conf,
                )
                cleaned = _clean_llm_json_response(response)
                return _parse_json_with_repair(cleaned)

            except json.JSONDecodeError as e:
                if attempt < _LLM_MAX_RETRIES:
                    logger.info(f"LLM JSON parse failed (attempt {attempt + 1}), retrying: {e}")
                    continue
                else:
                    logger.warning(f"LLM JSON parse failed (retries exhausted): {e}")
                    return None

    except Exception as e:
        logger.warning(f"LLM call failed: {e}")
        return None


def _normalize_for_comparison(text: str) -> str:
    """
    Normalize text for comparison (ref SmartResume's _normalize_for_comparison).

    Unify fullwidth/halfwidth, remove whitespace, Unicode normalization,
    so that "阿里巴巴" and "阿 里 巴 巴" can match.

    Args:
        text: Original text
    Returns:
        Normalized text
    """
    if not text:
        return ""
    # Unicode NFKC normalization (fullwidth to halfwidth, etc.)
    text = unicodedata.normalize("NFKC", text)
    # Remove all whitespace characters
    text = re.sub(r'\s+', '', text)
    return text.lower()

def _calc_single_exp_years(start_str: str, end_str: str) -> float:
    """
    Calculate years for a single experience entry.

    Args:
        start_str: Start date string
        end_str: End date string ("至今" etc. means current)
    Returns:
        Years (float, 1 decimal place), returns 0 if unable to calculate
    """
    from datetime import datetime

    start_str = str(start_str).strip()
    end_str = str(end_str).strip()
    if not start_str:
        return 0

    start_date = _parse_date_str(start_str)
    if not start_date:
        return 0

    if end_str in ("至今", "现在", "present", "Present", "now", "Now", ""):
        end_date = datetime.now()
    else:
        end_date = _parse_date_str(end_str)
        if not end_date:
            end_date = datetime.now()

    months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
    if months <= 0:
        return 0
    return round(months / 12.0, 1)


def _calculate_work_years(experiences: list[dict]) -> float:
    """
    Calculate total work years based on start/end dates of each work experience.

    Args:
        experiences: List of work experiences, each containing start_date, end_date fields
    Returns:
        Total work years (float), returns 0 if unable to calculate
    """
    total = 0.0
    for exp in experiences:
        total += _calc_single_exp_years(
            exp.get("start_date", ""), exp.get("end_date", "")
        )
    return round(total, 1)


def _parse_date_str(date_str: str) -> Optional[datetime.datetime]:
    """
    Parse date string, supporting multiple common formats.

    Supported formats:
    - 2024.1 / 2024.01
    - 2024-1 / 2024-01
    - 2024/1 / 2024/01
    - 2024年1月
    - 2024 (year only, defaults to January)

    Args:
        date_str: Date string
    Returns:
        datetime object, or None on parse failure
    """
    from datetime import datetime

    date_str = date_str.strip()
    # Try matching year.month / year-month / year/month / year(nian)month(yue) formats
    patterns = [
        (r"((?:19|20)\d{2})[.\-/年](\d{1,2})", "%Y-%m"),
        (r"^((?:19|20)\d{2})$", "%Y"),
    ]
    for pattern, _ in patterns:
        m = re.search(pattern, date_str)
        if m:
            try:
                year = int(m.group(1))
                month = int(m.group(2)) if len(m.groups()) > 1 else 1
                # Month range validation
                if month < 1 or month > 12:
                    month = 1
                return datetime(year, month, 1)
            except (ValueError, IndexError):
                continue
    return None




def _extract_description_from_range(
        index_range: list, lines: list[str],
        company: str = "", position: str = ""
) -> str:
    """
    Extract description from original text by index range (ref SmartResume's _extract_description_from_range).

    Key improvement:
    - Filter out lines containing both company name and position title (avoid mixing header lines into description)
    - Boundary safety checks

    Args:
        index_range: [start_line_number, end_line_number]
        lines: List of original line texts
        company: Company name (used to filter header lines)
        position: Position title (used to filter header lines)
    Returns:
        Extracted description text
    """
    if not index_range or len(index_range) != 2:
        return ""

    start_idx, end_idx = int(index_range[0]), int(index_range[1])

    # Boundary safety check
    if start_idx < 0 or end_idx >= len(lines) or start_idx > end_idx:
        return ""

    extracted_lines = lines[start_idx:end_idx + 1]

    # Filter out lines containing both company name and position title (ref SmartResume)
    if company or position:
        norm_company = _normalize_for_comparison(company)
        norm_position = _normalize_for_comparison(position)
        filtered = []
        for line in extracted_lines:
            norm_line = _normalize_for_comparison(line)
            # If a line contains both company name and position title, it's likely a header line, skip
            if norm_company and norm_position and norm_company in norm_line and norm_position in norm_line:
                continue
            # If a line exactly equals company name or position title, also skip
            if norm_line == norm_company or norm_line == norm_position:
                continue
            filtered.append(line)
        extracted_lines = filtered

    if not extracted_lines:
        return ""

    return "\n".join(line.strip() for line in extracted_lines if line.strip())


def _extract_basic_info(indexed_text: str, tenant_id , lang: str) -> Optional[dict]:
    """Extract basic info (subtask 1).

    Basic info is usually at the beginning of the resume, first 8000 chars suffice.
    """
    prompt = get_basic_info_prompt(lang).format(indexed_text=indexed_text[:8000])
    return _call_llm(prompt,tenant_id, lang)


def _extract_work_experience(indexed_text: str, tenant_id , lang: str) -> Optional[dict]:
    """Extract work experience (subtask 2, using index pointers).

    Work experience may span the middle-to-end of the resume, use full text to avoid truncation.
    """
    prompt = get_work_exp_prompt(lang).format(indexed_text=indexed_text)
    return _call_llm(prompt, tenant_id , lang)


def _extract_education(indexed_text: str, tenant_id , lang: str) -> Optional[dict]:
    """Extract education background (subtask 3).

    Education is usually at the end of the resume, must use full text to avoid truncation.
    Resume text is generally under 30K chars, within LLM context window.
    """
    prompt = get_education_prompt(lang).format(indexed_text=indexed_text)
    return _call_llm(prompt,tenant_id, lang)


def _extract_project_experience(indexed_text: str, tenant_id , lang: str) -> Optional[dict]:
    """Extract project experience (subtask 4, using index pointers).

    Project experience may span the middle-to-end of the resume, use full text to avoid truncation.
    """
    prompt = get_project_exp_prompt(lang).format(indexed_text=indexed_text)
    return _call_llm(prompt, tenant_id , lang)


def parse_with_llm(indexed_text: str, lines: list[str], tenant_id , lang: str) -> Optional[dict]:
    """
    Extract resume info using parallel task decomposition strategy (ref SmartResume Section 3.2).

    Decomposes extraction into four independent subtasks executed in parallel:
    1. Basic info (name, phone, skills, self-evaluation, etc.)
    2. Work experience (company, position, description line ranges)
    3. Education background (school, major, degree)
    4. Project experience (project name, role, description line ranges)

    Args:
        indexed_text: Line-indexed resume text
        lines: List of original line texts (for index-based extraction)
        lang: Language
    Returns:
        Merged structured resume dictionary, or None on failure
    """
    try:
        # Execute four subtasks in parallel
        with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
            future_basic = executor.submit(_extract_basic_info, indexed_text, tenant_id , lang)
            future_work = executor.submit(_extract_work_experience, indexed_text, tenant_id , lang)
            future_edu = executor.submit(_extract_education, indexed_text, tenant_id, lang)
            future_project = executor.submit(_extract_project_experience, indexed_text, tenant_id , lang)

            basic_info = future_basic.result(timeout=60)
            work_exp = future_work.result(timeout=60)
            education = future_edu.result(timeout=60)
            project_exp = future_project.result(timeout=60)

        # Merge results
        resume = {}

        # Merge basic info
        if basic_info:
            resume.update(basic_info)
            logger.info(f"Basic info extraction succeeded: {len(basic_info)} fields")

        # Process work experience (index pointer extraction)
        if work_exp and "workExperience" in work_exp:
            experiences = work_exp["workExperience"]
            companies = []
            positions = []
            work_descs = []
            # Save detailed info for each experience (dates, years) for chunk generation
            work_exp_details = []
            for exp in experiences:
                company = exp.get("company", "")
                position = exp.get("position", "")
                start_date = exp.get("start_date", "")
                end_date = exp.get("end_date", "")
                # Calculate years for this experience entry
                years = _calc_single_exp_years(start_date, end_date)
                if company:
                    companies.append(company)
                if position:
                    positions.append(position)
                # Save detailed info for each experience entry
                work_exp_details.append({
                    "company": company,
                    "position": position,
                    "start_date": start_date,
                    "end_date": end_date,
                    "years": years,
                })
                # Index pointer mechanism: extract description from original text by line range
                # Use _extract_description_from_range to filter header lines (ref SmartResume)
                desc_lines = exp.get("desc_lines", [])
                if isinstance(desc_lines, list) and len(desc_lines) == 2:
                    desc = _extract_description_from_range(
                        desc_lines, lines, company=company, position=position
                    )
                    if desc.strip():
                        work_descs.append(desc.strip())

            if companies:
                resume["corp_nm_tks"] = companies
                resume["corporation_name_tks"] = companies[0]
            if positions:
                resume["position_name_tks"] = positions
            if work_descs:
                resume["work_desc_tks"] = work_descs
            # Save experience details for _build_chunk_document
            if work_exp_details:
                resume["_work_exp_details"] = work_exp_details
            # Calculate total work years from each experience's dates (overrides LLM's guess in basic info)
            calculated_years = _calculate_work_years(experiences)
            if calculated_years > 0:
                resume["work_exp_flt"] = calculated_years
            logger.info(f"Work experience extraction succeeded: {len(experiences)} entries, calculated total years: {calculated_years}")

        # Process education background
        if education and "education" in education:
            edu_list = education["education"]
            schools = []
            majors = []
            degrees = []
            for edu in edu_list:
                if edu.get("school"):
                    schools.append(edu["school"])
                if edu.get("major"):
                    majors.append(edu["major"])
                if edu.get("degree"):
                    degrees.append(edu["degree"])
                # Extract graduation year
                end_date = edu.get("end_date", "")
                if end_date and not resume.get("edu_end_int"):
                    year_match = re.search(r"(19|20)\d{2}", str(end_date))
                    if year_match:
                        resume["edu_end_int"] = int(year_match.group(0))

            if schools:
                resume["school_name_tks"] = schools
                resume["first_school_name_tks"] = schools[-1]  # Earliest school is usually last
            if majors:
                resume["major_tks"] = majors
                resume["first_major_tks"] = majors[-1]
            if degrees:
                resume["degree_kwd"] = degrees
                # Infer highest degree (supports both Chinese and English degree names)
                degree_rank = {
                    "博士": 5, "PhD": 5, "Doctor": 5,
                    "硕士": 4, "Master": 4, "MBA": 4, "EMBA": 4, "MPA": 4,
                    "本科": 3, "Bachelor": 3,
                    "大专": 2, "专科": 2, "Associate": 2, "Diploma": 2,
                    "高中": 1, "High School": 1,
                }
                highest = max(degrees, key=lambda d: degree_rank.get(d, 0), default="")
                if highest:
                    resume["highest_degree_kwd"] = highest
                resume["first_degree_kwd"] = degrees[-1] if degrees else ""
            logger.info(f"Education extraction succeeded: {len(edu_list)} entries")

        # Process project experience (index pointer extraction, similar to work experience)
        if project_exp and "projectExperience" in project_exp:
            projects = project_exp["projectExperience"]
            project_names = []
            project_descs = []
            for proj in projects:
                name = proj.get("project_name", "")
                if name:
                    project_names.append(name)
                # Index pointer mechanism: extract project description from original text by line range
                desc_lines = proj.get("desc_lines", [])
                if isinstance(desc_lines, list) and len(desc_lines) == 2:
                    desc = _extract_description_from_range(
                        desc_lines, lines, company=name, position=proj.get("role", "")
                    )
                    if desc.strip():
                        project_descs.append(desc.strip())

            if project_names:
                resume["project_tks"] = project_names
            if project_descs:
                resume["project_desc_tks"] = project_descs
            logger.info(f"Project experience extraction succeeded: {len(projects)} entries")

        if not resume.get("name_kwd"):
            resume["name_kwd"] = "Unknown" if _is_english(lang) else "未知"

        return resume if len(resume) > 2 else None

    except concurrent.futures.TimeoutError:
        logger.warning("LLM parallel extraction timed out")
        return None
    except Exception as e:
        logger.warning(f"LLM parallel extraction failed: {e}")
        return None


# ==================== Phase 3: Regex Fallback Parsing ====================



def parse_with_regex(text: str, lang: str = "Chinese") -> dict:
    """
    Parse resume text using regex (fallback strategy)

    When LLM parsing fails, use regex to extract basic structured info from text.

    Args:
        text: Resume text content (without line number index)
        lang: Language parameter, default "Chinese"
    Returns:
        Structured resume info dictionary
    """
    resume: dict = {}
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # --- Extract Name ---
    if _is_english(lang):
        # English resume: extract from "Name: XXX" format
        for line in lines[:30]:
            name_match = re.search(r'(?:Name|Full\s*Name)\s*[:：]\s*([A-Za-z][A-Za-z\s\-\.]{1,40})', line, re.IGNORECASE)
            if name_match:
                resume["name_kwd"] = name_match.group(1).strip()
                break
        # English resume strategy 2: first line if short text without digits, may be a name
        if "name_kwd" not in resume and lines:
            first = lines[0].strip()
            if len(first) <= 40 and not re.search(r"\d", first) and re.match(r'^[A-Za-z][A-Za-z\s\-\.]+$', first):
                resume["name_kwd"] = first
    else:
        # Chinese resume: extract from "姓名：XXX" format
        for line in lines[:30]:
            name_match = re.search(r'姓\s*名\s*[:：]\s*([\u4e00-\u9fa5]{2,4})', line)
            if name_match:
                resume["name_kwd"] = name_match.group(1)
                break

        # Strategy 2: search first 20 lines for standalone Chinese names (2-4 chars), excluding common title words
        if "name_kwd" not in resume:
            title_words = {
                "个人", "简历", "求职", "应聘", "基本", "信息", "概述", "简介",
                "教育", "工作", "经历", "经验", "技能", "项目", "自我", "评价",
                "专业", "技术", "证书", "语言", "能力", "培训", "荣誉", "奖项",
            }
            for line in lines[:20]:
                if any(w in line for w in title_words):
                    continue
                if re.search(r'[:：]', line) and len(line) > 6:
                    continue
                cleaned = re.sub(r"^[A-Za-z_\-\d\s]+\s+", "", line)
                cleaned = re.sub(r"\s+[A-Za-z_\-\d\s]+$", "", cleaned).strip()
                if 2 <= len(cleaned) <= 4 and re.match(r"^[\u4e00-\u9fa5]{2,4}$", cleaned):
                    resume["name_kwd"] = cleaned
                    break

        # Strategy 3: first line if short without digits, may be a name
        if "name_kwd" not in resume and lines:
            first = lines[0].strip()
            if len(first) <= 10 and not re.search(r"\d", first):
                cn_part = re.findall(r'[\u4e00-\u9fa5]+', first)
                if cn_part and 2 <= len(cn_part[0]) <= 4:
                    resume["name_kwd"] = cn_part[0]

    # --- Extract Phone Number ---
    phones = re.findall(r"1[3-9]\d{9}", text)
    if phones:
        resume["phone_kwd"] = phones[0]

    # --- Extract Email ---
    emails = re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    if emails:
        resume["email_tks"] = emails[0]

    # --- Extract Gender ---
    if _is_english(lang):
        # English resume: extract from "Gender: Male/Female" format
        gender_label = re.search(r'(?:Gender|Sex)\s*[:：]\s*(Male|Female|M|F)', text, re.IGNORECASE)
        if gender_label:
            raw = gender_label.group(1).strip().upper()
            resume["gender_kwd"] = "Male" if raw in ("M", "MALE") else "Female"
        else:
            gender_match = re.search(r'\b(Male|Female)\b', text[:500], re.IGNORECASE)
            if gender_match:
                resume["gender_kwd"] = gender_match.group(1).capitalize()
    else:
        # Chinese resume: extract from "性别：男/女" format
        gender_label = re.search(r'性\s*别\s*[:：]\s*(男|女)', text)
        if gender_label:
            resume["gender_kwd"] = gender_label.group(1)
        else:
            gender_match = re.search(r"(男|女)", text[:500])
            if gender_match:
                resume["gender_kwd"] = gender_match.group(1)

    # --- Extract Age ---
    if _is_english(lang):
        # English resume: match "25 years old" or "Age: 25"
        age_match = re.search(r'(?:Age)\s*[:：]\s*(\d{1,2})', text, re.IGNORECASE)
        if not age_match:
            age_match = re.search(r'(\d{1,2})\s*years?\s*old', text, re.IGNORECASE)
        if age_match:
            resume["age_int"] = int(age_match.group(1))
    else:
        # Chinese resume: match "25岁"
        age_match = re.search(r"(\d{1,2})\s*岁", text)
        if age_match:
            resume["age_int"] = int(age_match.group(1))

    # --- Extract Date of Birth ---
    if _is_english(lang):
        # English resume: match "1990-01-15" or "Jan 15, 1990" etc.
        birth_match = re.search(r'(?:Birth|DOB|Date\s*of\s*Birth)\s*[:：]\s*(.{6,20})', text, re.IGNORECASE)
        if birth_match:
            resume["birth_dt"] = birth_match.group(1).strip()
        else:
            birth_match = re.search(r"(19|20)\d{2}[-/]\d{1,2}[-/]\d{1,2}", text)
            if birth_match:
                resume["birth_dt"] = birth_match.group(0)
    else:
        # Chinese resume: match "1990年1月15日" or "1990-01-15"
        birth_match = re.search(r"(19|20)\d{2}[年/-]\d{1,2}[月/-]\d{1,2}", text)
        if birth_match:
            resume["birth_dt"] = birth_match.group(0)

    # --- Extract Education Level ---
    degree_keywords_zh = ["博士", "硕士", "本科", "大专", "专科", "高中", "MBA", "EMBA", "MPA"]
    degree_keywords_en = ["PhD", "Master", "Bachelor", "Associate", "Diploma", "High School",
                          "MBA", "EMBA", "MPA", "Doctor"]
    degree_keywords = degree_keywords_en if _is_english(lang) else degree_keywords_zh
    found_degrees = [d for d in degree_keywords if d in text]
    if found_degrees:
        resume["degree_kwd"] = found_degrees

    # --- Extract School ---
    if _is_english(lang):
        # English resume: match "University/College/Institute/School" keywords
        schools = re.findall(
            r'([A-Z][A-Za-z\s\-&]{2,40}(?:University|College|Institute|School|Academy))',
            text
        )
        # Remove extra whitespace
        schools = [re.sub(r'\s+', ' ', s).strip() for s in schools]
    else:
        # Chinese resume: match "XX大学/学院/职业技术学院"
        schools = re.findall(r"[\u4e00-\u9fa5]{2,15}(?:大学|学院|职业技术学院)", text)
    if schools:
        resume["school_name_tks"] = list(set(schools))
        resume["first_school_name_tks"] = schools[0]

    # --- Extract Major ---
    if _is_english(lang):
        # English resume: match "Major: XXX" / "Field of Study: XXX" / "Specialization: XXX"
        majors = re.findall(
            r'(?:Major|Field\s*of\s*Study|Specialization|Concentration)\s*[:：]\s*([A-Za-z\s\-&,]{2,40})',
            text, re.IGNORECASE
        )
        majors = [m.strip() for m in majors if m.strip()]
    else:
        # Chinese resume: match "专业：XXX"
        majors = re.findall(r"专业[:：]\s*([\u4e00-\u9fa5]{2,20})", text)
    if majors:
        resume["major_tks"] = majors
        resume["first_major_tks"] = majors[0]

    # --- Extract Company Names ---
    if _is_english(lang):
        # English resume: match common company suffixes
        en_company_patterns = [
            r'([A-Z][A-Za-z\s\-&,\.]{2,40}(?:Inc\.|Corp\.|Ltd\.|LLC|Co\.|Company|Group|Technologies|Technology|Solutions|Consulting|Services|Bank))',
        ]
        companies = []
        for pattern in en_company_patterns:
            companies.extend(re.findall(pattern, text))
        companies = [re.sub(r'\s+', ' ', c).strip() for c in companies]
    else:
        # Chinese resume: match "XX有限公司" format
        company_patterns = [
            r"[\u4e00-\u9fa5]{2,20}[（(][\u4e00-\u9fa5]{2,10}[)）](?:科技|信息技术|网络科技)?(?:股份)?有限公司",
            r"[\u4e00-\u9fa5]{4,20}(?:科技|信息技术|网络科技|银行)?(?:股份)?有限公司",
        ]
        companies = []
        for pattern in company_patterns:
            companies.extend(re.findall(pattern, text))

    unique_companies = []
    seen = set()
    # Filter verb list (bilingual)
    filter_verbs = (
        ["completed", "conducted", "implemented", "responsible", "participated", "developed"]
        if _is_english(lang)
        else ["完成", "进行", "实施", "负责", "参与", "开发"]
    )
    min_len = 3 if _is_english(lang) else 6
    for c in companies:
        if len(c) < min_len or any(v in c.lower() for v in filter_verbs) or c in seen:
            continue
        is_sub = False
        for existing in list(unique_companies):
            if c in existing:
                is_sub = True
                break
            if existing in c:
                unique_companies.remove(existing)
                seen.discard(existing)
        if not is_sub:
            unique_companies.append(c)
            seen.add(c)

    if unique_companies:
        resume["corp_nm_tks"] = unique_companies
        resume["corporation_name_tks"] = unique_companies[0]

    # --- Extract Position (improved: context constraints to reduce noise) ---
    if _is_english(lang):
        # English resume: Strategy 1 - extract from "Title: XXX" / "Position: XXX" / "Role: XXX" format
        position_label_matches = re.findall(
            r'(?:Title|Position|Role|Job\s*Title)\s*[:：]\s*([A-Za-z\s\-/&]{2,30})',
            text, re.IGNORECASE
        )
        positions = [p.strip() for p in position_label_matches if p.strip()]

        # English resume: Strategy 2 - match common position suffix keywords
        en_position_suffixes = [
            "Engineer", "Manager", "Director", "Supervisor", "Specialist",
            "Designer", "Consultant", "Assistant", "Architect", "Analyst",
            "Developer", "Lead", "Officer", "Coordinator", "Administrator",
            "Intern", "VP", "President",
        ]
        for line in lines:
            if len(line) > 60:
                continue  # Skip overly long lines (usually description text)
            for suffix in en_position_suffixes:
                match = re.search(rf'([A-Za-z\s\-]{{1,25}}{suffix})\b', line, re.IGNORECASE)
                if match:
                    pos = match.group(1).strip()
                    # Filter out matches that are clearly not positions (contain verbs)
                    filter_pos_verbs = ["responsible", "participated", "completed", "developed", "designed"]
                    if not any(v in pos.lower() for v in filter_pos_verbs) and len(pos) > 3:
                        positions.append(pos)
    else:
        # Chinese resume: Strategy 1 - extract from "职位/岗位：XXX" format
        position_label_matches = re.findall(
            r'(?:职位|岗位|职务|职称|担任)\s*[:：]\s*([\u4e00-\u9fa5a-zA-Z]{2,15})',
            text
        )
        positions = list(position_label_matches)

        # Chinese resume: Strategy 2 - extract from work experience paragraphs (company name followed by position)
        for line in lines:
            pos_match = re.search(
                r'(?:有限公司|集团|银行)\s+([\u4e00-\u9fa5]{2,8}(?:工程师|经理|总监|主管|专员|设计师|顾问|助理|架构师|分析师|运营|产品))',
                line
            )
            if pos_match:
                positions.append(pos_match.group(1))

        # Chinese resume: Strategy 3 - position keywords in standalone lines (length-limited to avoid matching description text)
        position_suffixes = ["工程师", "经理", "总监", "主管", "专员", "设计师", "顾问",
                             "助理", "架构师", "分析师", "开发者", "负责人"]
        for line in lines:
            if len(line) > 20:
                continue  # Skip overly long lines
            for suffix in position_suffixes:
                match = re.search(rf'([\u4e00-\u9fa5]{{1,6}}{suffix})', line)
                if match:
                    pos = match.group(1)
                    if not any(v in pos for v in ["负责", "参与", "完成", "开发了", "设计了"]):
                        positions.append(pos)

    if positions:
        # Deduplicate while preserving order
        seen_pos = set()
        unique_positions = []
        for p in positions:
            if p not in seen_pos:
                seen_pos.add(p)
                unique_positions.append(p)
        resume["position_name_tks"] = unique_positions

    # --- Extract Years of Experience ---
    if _is_english(lang):
        # English resume: match "5 years experience" / "5+ years of experience"
        work_exp_match = re.search(r'(\d+)\+?\s*years?\s*(?:of\s*)?(?:experience|work)', text, re.IGNORECASE)
        if work_exp_match:
            resume["work_exp_flt"] = float(work_exp_match.group(1))
    else:
        # Chinese resume: match "5年...经验"
        work_exp_match = re.search(r"(\d+)\s*年.*?经验", text)
        if work_exp_match:
            resume["work_exp_flt"] = float(work_exp_match.group(1))

    # --- Extract Graduation Year ---
    if _is_english(lang):
        # English resume: match "Graduated 2020" / "Graduation: 2020" / "Class of 2020"
        grad_match = re.search(r'(?:Graduat(?:ed|ion)|Class\s*of)\s*[:：]?\s*((?:19|20)\d{2})', text, re.IGNORECASE)
        if grad_match:
            resume["edu_end_int"] = int(grad_match.group(1))
    else:
        # Chinese resume: match "2020年...毕业"
        grad_match = re.search(r"((?:19|20)\d{2})\s*年.*?毕业", text)
        if grad_match:
            resume["edu_end_int"] = int(grad_match.group(1))

    if "name_kwd" not in resume:
        resume["name_kwd"] = "Unknown" if _is_english(lang) else "未知"

    return resume



# ==================== Phase 4: Post-processing Pipeline ====================


def _postprocess_resume(resume: dict, lines: list[str], lang: str = "Chinese") -> dict:
    """
    Four-phase post-processing pipeline (ref: SmartResume Section 3.2.3)

    1. Source text validation: check if key fields can be found in the original text
    2. Domain normalization: standardize date formats, clean company name suffix noise
    3. Contextual deduplication: remove duplicate company/school entries
    4. Field completion: ensure all required fields exist

    Args:
        resume: Raw resume dictionary extracted by LLM
        lines: Original line text list (for source text validation)
        lang: Language parameter, default "Chinese"
    Returns:
        Post-processed resume dictionary
    """
    _en = _is_english(lang)
    full_text = "\n".join(lines) if lines else ""
    # Normalize full text for comparison (ref: SmartResume _validate_fields_in_text)
    norm_full_text = _normalize_for_comparison(full_text)

    # --- Phase 1: Source text validation (prune hallucinations, ref: SmartResume _validate_fields_in_text) ---
    # Name validation: clear if not found in source text (SmartResume strategy: discard hallucinated fields)
    _unknown_names = ("未知", "Unknown")
    if resume.get("name_kwd") and resume["name_kwd"] not in _unknown_names:
        norm_name = _normalize_for_comparison(resume["name_kwd"])
        if norm_full_text and norm_name and norm_name not in norm_full_text:
            logger.warning(f"Name '{resume['name_kwd']}' not found in source text, classified as LLM hallucination, cleared")
            resume["name_kwd"] = ""

    # Validate company names (strict matching: full name must appear in source text, no longer using loose 4-char prefix matching)
    if resume.get("corp_nm_tks") and norm_full_text:
        verified_companies = []
        for company in resume["corp_nm_tks"]:
            norm_company = _normalize_for_comparison(company)
            if norm_company and norm_company in norm_full_text:
                verified_companies.append(company)
            else:
                logger.debug(f"Company '{company}' not found in source text, filtered out")
        # Update even if all filtered out (SmartResume strategy: prefer missing over wrong)
        resume["corp_nm_tks"] = verified_companies
        if verified_companies:
            resume["corporation_name_tks"] = verified_companies[0]
        else:
            resume["corporation_name_tks"] = ""

    # Validate school names (ref: SmartResume _validate_fields_in_text)
    if resume.get("school_name_tks") and norm_full_text:
        verified_schools = []
        for school in resume["school_name_tks"]:
            norm_school = _normalize_for_comparison(school)
            if norm_school and norm_school in norm_full_text:
                verified_schools.append(school)
            else:
                logger.debug(f"School '{school}' not found in source text, filtered out")
        resume["school_name_tks"] = verified_schools
        if verified_schools:
            if resume.get("first_school_name_tks"):
                # Ensure first_school is also in the verified list
                if resume["first_school_name_tks"] not in verified_schools:
                    resume["first_school_name_tks"] = verified_schools[-1]
        else:
            resume["first_school_name_tks"] = ""

    # Validate position names
    if resume.get("position_name_tks") and norm_full_text:
        verified_positions = []
        for pos in resume["position_name_tks"]:
            norm_pos = _normalize_for_comparison(pos)
            if norm_pos and norm_pos in norm_full_text:
                verified_positions.append(pos)
        if verified_positions:
            resume["position_name_tks"] = verified_positions

    # --- Phase 2: Domain normalization ---
    # Standardize date format
    if resume.get("birth_dt"):
        resume["birth_dt"] = re.sub(r"[年月]", "-", str(resume["birth_dt"])).rstrip("-")

    # Clean non-digit characters from phone number (keep + sign)
    if resume.get("phone_kwd"):
        phone = re.sub(r"[^\d+]", "", str(resume["phone_kwd"]))
        if phone:
            resume["phone_kwd"] = phone

    # Standardize gender (output format determined by language parameter)
    if resume.get("gender_kwd"):
        gender = str(resume["gender_kwd"]).strip()
        if gender in ("male", "Male", "M", "m", "男"):
            resume["gender_kwd"] = "Male" if _en else "男"
        elif gender in ("female", "Female", "F", "f", "女"):
            resume["gender_kwd"] = "Female" if _en else "女"

    # --- Phase 3: Contextual deduplication ---
    for list_field in ["corp_nm_tks", "school_name_tks", "major_tks",
                       "position_name_tks", "skill_tks"]:
        if isinstance(resume.get(list_field), list):
            # Order-preserving deduplication
            seen = set()
            deduped = []
            for item in resume[list_field]:
                item_str = str(item).strip()
                if item_str and item_str not in seen:
                    seen.add(item_str)
                    deduped.append(item_str)
            resume[list_field] = deduped

    # --- Phase 4: Field completion ---
    required_fields = [
        "name_kwd", "gender_kwd", "phone_kwd", "email_tks",
        "position_name_tks", "school_name_tks", "major_tks",
    ]
    for field in required_fields:
        if field not in resume:
            if field.endswith("_tks"):
                resume[field] = []
            elif field.endswith("_int") or field.endswith("_flt"):
                resume[field] = 0
            else:
                resume[field] = ""

    # Clean internal marker fields (already handled in Phase 1, this is a safety fallback)
    resume.pop("_name_confidence", None)

    return resume


# ==================== Pipeline Orchestration & Chunk Construction ====================


def parse_resume(filename: str, binary: bytes, tenant_id , lang: str = "Chinese") -> tuple[dict, list[str], list[dict]]:
    """
    Resume parsing pipeline orchestration function

    Execution flow:
        1. Text extraction (dual-path fusion + layout reconstruction + line-number index)
        2. Parallel LLM structured extraction (three sub-tasks)
        3. Regex fallback parsing (when LLM fails)
        4. Four-phase post-processing

    Args:
        filename: File name
        binary: File binary content
        lang: Language, default "Chinese"
    Returns:
        (resume, lines, line_positions) tuple:
        - resume: Structured resume information dictionary
        - lines: Original line text list (for chunk text matching and positioning)
        - line_positions: Per-line coordinate info list (for writing chunk position_int fields)
    """
    # Phase 1: Text extraction
    indexed_text, lines, line_positions = extract_text(filename, binary)
    if not indexed_text or not lines:
        logger.warning(f"Text extraction returned empty: {filename}")
        default_name = "Unknown" if _is_english(lang) else "未知"
        return {"name_kwd": default_name}, [], []

    # Phase 2: Parallel LLM structured extraction
    resume = parse_with_llm(indexed_text, lines, tenant_id , lang)

    # Phase 3: Fallback to regex parsing when LLM fails
    if not resume:
        logger.info(f"LLM parsing failed, falling back to regex parsing: {filename}")
        plain_text = "\n".join(lines)
        resume = parse_with_regex(plain_text, lang)

    # Phase 4: Post-processing pipeline
    resume = _postprocess_resume(resume, lines, lang)

    return resume, lines, line_positions


def _build_chunk_document(filename: str, resume: dict,
                          lang: str = "Chinese") -> list[dict]:
    """
    Build a list of document chunks from structured resume information

    Each field generates an independent chunk containing tokenization results and metadata.
    Compatible with the build_chunks flow in task_executor.py.

    Key design: Each chunk redundantly includes key identity fields (name, phone, email, etc.),
    so that when any chunk is retrieved, the candidate's identity can be immediately identified.
    The full resume can be fetched via doc_id to get all chunks for complete information.

    Args:
        filename: File name
        resume: Structured resume information dictionary
        lang: Language parameter, default "Chinese"
    Returns:
        Document chunk list, each chunk contains content_with_weight, content_ltks,
        position_int, page_num_int, top_int and other fields
    """
    chunks = []
    # Get the corresponding field map version based on language parameter
    field_map = get_field_map(lang)
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename)),
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])

    # Extract key identity fields, redundantly written to each chunk
    # These fields are small in size but high in information density; once retrieved, the candidate can be immediately identified
    _IDENTITY_FIELDS = ("name_kwd", "phone_kwd", "email_tks", "gender_kwd",
                        "highest_degree_kwd", "work_exp_flt", "corporation_name_tks")
    identity_meta = {}
    for ik in _IDENTITY_FIELDS:
        iv = resume.get(ik)
        if not iv:
            continue
        if ik.endswith("_tks"):
            identity_meta[ik] = rag_tokenizer.tokenize(
                " ".join(iv) if isinstance(iv, list) else str(iv)
            )
        elif ik.endswith("_kwd"):
            identity_meta[ik] = iv if isinstance(iv, list) else str(iv)
        elif ik.endswith("_flt"):
            try:
                identity_meta[ik] = float(iv)
            except (ValueError, TypeError):
                pass
        else:
            identity_meta[ik] = str(iv)

    # Build resume summary text, appended to each chunk's content to improve semantic retrieval recall
    summary_parts = []
    _en = _is_english(lang)
    if resume.get("name_kwd"):
        summary_parts.append(f"{'Name' if _en else '姓名'}:{resume['name_kwd']}")
    if resume.get("phone_kwd"):
        summary_parts.append(f"{'Phone' if _en else '电话'}:{resume['phone_kwd']}")
    if resume.get("corporation_name_tks"):
        corp = resume["corporation_name_tks"]
        summary_parts.append(f"{'Company' if _en else '公司'}:{corp if isinstance(corp, str) else ' '.join(corp)}")
    if resume.get("highest_degree_kwd"):
        summary_parts.append(f"{'Degree' if _en else '学历'}:{resume['highest_degree_kwd']}")
    if resume.get("work_exp_flt"):
        if _en:
            summary_parts.append(f"Experience:{resume['work_exp_flt']}yrs")
        else:
            summary_parts.append(f"经验:{resume['work_exp_flt']}年")
    resume_summary = " | ".join(summary_parts) if summary_parts else ""

    # List fields that need per-element splitting (each experience/project generates a separate chunk to avoid oversized merged chunks)
    _SPLIT_LIST_FIELDS = {"work_desc_tks", "project_desc_tks"}

    # Basic info field set: these fields should be merged into one chunk to avoid splitting name, phone, email, etc.
    _BASIC_INFO_FIELDS = {
        "name_kwd", "name_pinyin_kwd", "gender_kwd", "age_int",
        "phone_kwd", "email_tks", "birth_dt", "work_exp_flt",
        "position_name_tks", "expect_city_names_tks",
        "expect_position_name_tks",
    }

    # Education field set: degree, school, major, tags, etc. should be merged into one chunk
    _EDUCATION_FIELDS = {
        "first_school_name_tks", "first_degree_kwd", "highest_degree_kwd",
        "first_major_tks", "edu_first_fea_kwd", "degree_kwd", "major_tks",
        "school_name_tks", "sch_rank_kwd", "edu_fea_kwd", "edu_end_int",
    }

    # Skills & certificates field set: skills, languages, certificates are small, merge into one chunk
    _SKILL_CERT_FIELDS = {
        "skill_tks", "language_tks", "certificate_tks",
    }

    # Work overview field set: company list, industry, most recent company merged into one chunk
    _WORK_OVERVIEW_FIELDS = {
        "corporation_name_tks", "corp_nm_tks", "industry_name_tks",
    }

    # All merge groups: (field_set, group_title) tuple list
    _MERGE_GROUPS = [
        (_BASIC_INFO_FIELDS, "Basic Info" if _en else "基本信息"),
        (_EDUCATION_FIELDS, "Education" if _en else "教育背景"),
        (_SKILL_CERT_FIELDS, "Skills & Certificates" if _en else "技能与证书"),
        (_WORK_OVERVIEW_FIELDS, "Work Overview" if _en else "工作概况"),
    ]

    # Collect all fields that need merge processing; skip them during individual iteration
    _ALL_MERGED_FIELDS = set()
    for fields_set, _ in _MERGE_GROUPS:
        _ALL_MERGED_FIELDS.update(fields_set)

    # Merge fields by group, generating one chunk per group
    for fields_set, group_title in _MERGE_GROUPS:
        group_parts = []
        group_field_values = {}  # Store structured values for each field, to be written into chunk
        for field_key in field_map:
            if field_key not in fields_set:
                continue
            value = resume.get(field_key)
            if not value:
                continue
            field_desc = field_map[field_key]
            if isinstance(value, list):
                text_value = " ".join(str(v) for v in value if v)
            else:
                text_value = str(value)
            if not text_value.strip():
                continue
            group_parts.append(f"{field_desc}: {text_value}")
            group_field_values[field_key] = value

        if not group_parts:
            continue

        content = f"{group_title}\n" + "\n".join(group_parts)
        if resume_summary:
            content += f"\n[{resume_summary}]"
        chunk = {
            "content_with_weight": content,
            "content_ltks": rag_tokenizer.tokenize(content),
            "content_sm_ltks": rag_tokenizer.fine_grained_tokenize(
                rag_tokenizer.tokenize(content)
            ),
        }
        chunk.update(doc)
        # Redundantly write identity fields
        for mk, mv in identity_meta.items():
            chunk[mk] = mv
        # Write each field's structured value into chunk (for structured retrieval)
        for fk, fv in group_field_values.items():
            if fk.endswith("_tks"):
                text_val = " ".join(str(v) for v in fv) if isinstance(fv, list) else str(fv)
                chunk[fk] = rag_tokenizer.tokenize(text_val)
            elif fk.endswith("_kwd"):
                chunk[fk] = fv if isinstance(fv, list) else str(fv)
            elif fk.endswith("_int"):
                try:
                    chunk[fk] = int(fv)
                except (ValueError, TypeError):
                    pass
            elif fk.endswith("_flt"):
                try:
                    chunk[fk] = float(fv)
                except (ValueError, TypeError):
                    pass
            else:
                chunk[fk] = str(fv)
        chunks.append(chunk)

    # Iterate over field map, generating a chunk for each non-merged field with a value
    for field_key, field_desc in field_map.items():
        # Skip fields already processed in merge groups
        if field_key in _ALL_MERGED_FIELDS:
            continue
        value = resume.get(field_key)
        if not value:
            continue

        # For work/project descriptions (long text lists), split into multiple chunks per element
        if field_key in _SPLIT_LIST_FIELDS and isinstance(value, list):
            # Get company name list to add context to each work description
            corp_list = resume.get("corp_nm_tks", []) if field_key == "work_desc_tks" else []
            project_list = resume.get("project_tks", []) if field_key == "project_desc_tks" else []
            # Get detailed info for each work experience entry (time period, years)
            work_details = resume.get("_work_exp_details", []) if field_key == "work_desc_tks" else []

            for idx, item in enumerate(value):
                item_text = str(item).strip()
                if not item_text:
                    continue

                # Add company/project name prefix to each description for context
                if field_key == "work_desc_tks" and idx < len(work_details):
                    # Use detailed info to build prefix, including company, time range, years
                    detail = work_details[idx]
                    company = detail.get("company", "")
                    start_d = detail.get("start_date", "")
                    end_d = detail.get("end_date", "")
                    years = detail.get("years", 0)
                    # Build time range text
                    time_parts = []
                    if start_d:
                        time_range = f"{start_d}-{end_d}" if end_d else str(start_d)
                        time_parts.append(time_range)
                    if years > 0:
                        time_parts.append(f"{years}{'yrs' if _en else '年'}")
                    time_text = " ".join(time_parts)
                    if company and time_text:
                        content_prefix = f"{field_desc}（{company} {time_text}）"
                    elif company:
                        content_prefix = f"{field_desc}（{company}）"
                    else:
                        content_prefix = f"{field_desc}（{'#' if _en else '第'}{idx + 1}{'' if _en else '段'}）"
                elif field_key == "work_desc_tks" and idx < len(corp_list):
                    content_prefix = f"{field_desc}（{corp_list[idx]}）"
                elif field_key == "project_desc_tks" and idx < len(project_list):
                    content_prefix = f"{field_desc}（{project_list[idx]}）"
                else:
                    content_prefix = f"{field_desc}（{'#' if _en else '第'}{idx + 1}{'' if _en else '段'}）"

                if resume_summary:
                    content = f"{content_prefix}: {item_text}\n[{resume_summary}]"
                else:
                    content = f"{content_prefix}: {item_text}"

                chunk = {
                    "content_with_weight": content,
                    "content_ltks": rag_tokenizer.tokenize(content),
                    "content_sm_ltks": rag_tokenizer.fine_grained_tokenize(
                        rag_tokenizer.tokenize(content)
                    ),
                }
                chunk.update(doc)

                # Redundantly write identity fields
                for mk, mv in identity_meta.items():
                    if mk != field_key:
                        chunk[mk] = mv

                # Tokenization result for current segment
                chunk[field_key] = rag_tokenizer.tokenize(item_text)
                chunks.append(chunk)
            continue

        # Merge list values into text
        if isinstance(value, list):
            text_value = " ".join(str(v) for v in value if v)
        else:
            text_value = str(value)

        if not text_value.strip():
            continue

        # Build chunk content: "field_desc: field_value", append summary for semantic association
        if resume_summary and field_key not in ("name_kwd", "phone_kwd"):
            content = f"{field_desc}: {text_value}\n[{resume_summary}]"
        else:
            content = f"{field_desc}: {text_value}"

        chunk = {
            "content_with_weight": content,
            "content_ltks": rag_tokenizer.tokenize(content),
            "content_sm_ltks": rag_tokenizer.fine_grained_tokenize(
                rag_tokenizer.tokenize(content)
            ),
        }
        chunk.update(doc)

        # Redundantly write identity fields (do not overwrite the current field's own value)
        for mk, mv in identity_meta.items():
            if mk != field_key:
                chunk[mk] = mv

        # Write resume field value into the chunk's corresponding field (for structured retrieval)
        if field_key.endswith("_tks"):
            chunk[field_key] = rag_tokenizer.tokenize(text_value)
        elif field_key.endswith("_kwd"):
            if isinstance(value, list):
                chunk[field_key] = value
            else:
                chunk[field_key] = text_value
        elif field_key.endswith("_int"):
            try:
                chunk[field_key] = int(value)
            except (ValueError, TypeError):
                pass
        elif field_key.endswith("_flt"):
            try:
                chunk[field_key] = float(value)
            except (ValueError, TypeError):
                pass
        else:
            chunk[field_key] = text_value

        chunks.append(chunk)

    # If no chunks were generated, create at least one chunk containing the name
    if not chunks:
        name = resume.get("name_kwd", "Unknown" if _en else "未知")
        content = f"{'Name' if _en else '姓名'}: {name}"
        chunk = {
            "content_with_weight": content,
            "content_ltks": rag_tokenizer.tokenize(content),
            "content_sm_ltks": rag_tokenizer.fine_grained_tokenize(
                rag_tokenizer.tokenize(content)
            ),
        }
        chunk.update(doc)
        chunks.append(chunk)

    # Write coordinate info to each chunk (position_int, page_num_int, top_int)
    #
    # Resume chunks are split by semantic fields (basic info, education, work description, etc.),
    # not by PDF physical regions. Field values may be scattered across multiple locations in the PDF,
    # and using text matching to reverse-lookup coordinates would cause disordered sorting.
    #
    # Therefore, assign incrementing coordinates based on chunk generation order (i.e., semantic logical order),
    # ensuring display order: basic info -> education -> skills/certs -> work overview -> work desc -> project desc...
    #
    # add_positions input format: [(page, left, right, top, bottom), ...]
    #   - page starts from 0, function internally stores +1
    #   - task_executor sorts by page_num_int and top_int (page first, then Y coordinate)
    from rag.nlp import add_positions

    for i, ck in enumerate(chunks):
        # All chunks placed on page=0, top increments by index to ensure logical ordering
        add_positions(ck, [[0, 0, 0, i, i]])

    return chunks

def _blackout_text_regions(image: "np.ndarray", meta_blocks: list[dict], page_idx: int,
                           pdf_to_img_scale: float) -> "np.ndarray":
    """
    Black out metadata-extracted text regions on the page image to prevent OCR duplication.

    Ref: SmartResume blackout strategy — extract metadata text first, black out those regions,
    then run OCR on the blacked-out image so it only recognizes content metadata missed.
    More reliable than IoU-based deduplication.

    Args:
        image: Page image (numpy array)
        meta_blocks: Text blocks from metadata extraction
        page_idx: Current page number
        pdf_to_img_scale: Scale factor from PDF coordinates to image coordinates
    Returns:
        Image with text regions blacked out
    """
    import cv2
    blacked = image.copy()
    page_blocks = [b for b in meta_blocks if b.get("page") == page_idx]
    # Draw filled black rectangles over each metadata text block
    padding = 2  # Extra pixels to ensure full coverage
    for b in page_blocks:
        x0 = int(b["x0"] * pdf_to_img_scale) - padding
        y0 = int(b["top"] * pdf_to_img_scale) - padding
        x1 = int(b["x1"] * pdf_to_img_scale) + padding
        y1 = int(b["bottom"] * pdf_to_img_scale) + padding
        # Clamp to image boundaries
        x0 = max(0, x0)
        y0 = max(0, y0)
        x1 = min(blacked.shape[1], x1)
        y1 = min(blacked.shape[0], y1)
        cv2.rectangle(blacked, (x0, y0), (x1, y1), (0, 0, 0), -1)
    return blacked



def chunk(filename, binary, tenant_id, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """
    Resume parsing entry function (compatible with task_executor.py)

    This function is the entry point registered as FACTORY[ParserType.RESUME.value],
    with a signature consistent with other parsers (e.g., naive.chunk).

    Args:
        filename: File name
        binary: File binary content
        from_page: Start page number (not used in resume parsing)
        to_page: End page number (not used in resume parsing)
        lang: Language, default "Chinese"
        callback: Progress callback function, accepts (progress, message) parameters
        **kwargs: Other parameters (parser_config, kb_id, tenant_id, etc.)
    Returns:
        Document chunk list
    """
    if callback is None:
        def callback(prog, msg): return None

    try:
        callback(0.1, "Starting resume parsing...")

        # Parse resume
        resume, lines, line_positions = parse_resume(filename, binary, tenant_id , lang)
        callback(0.6, "Resume structured extraction complete")

        # Build document chunks (with coordinate info)
        chunks = _build_chunk_document(filename, resume, lang)
        callback(0.9, f"Document chunk construction complete, {len(chunks)} chunks total")

        callback(1.0, "Resume parsing complete")
        return chunks

    except Exception as e:
        logger.exception(f"Resume parsing exception: {filename}")
        callback(-1, f"Resume parsing failed: {str(e)}")
        return []