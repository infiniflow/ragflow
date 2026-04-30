#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#
import re
from io import BytesIO

from bs4 import BeautifulSoup
from docx import Document
from api.db.services.llm_service import LLMBundle
from api.db.joint_services.tenant_model_service import (
    get_model_config_by_type_and_name,
    get_tenant_default_model_by_type,
)
from common.constants import LLMType
from deepdoc.parser.figure_parser import VisionFigureParser
from rag.nlp import is_english, random_choices, remove_contents_table


def remove_toc(items):
    indexed = [(_item_text(item), i) for i, item in enumerate(items)]
    remove_contents_table(indexed, eng=_is_english(indexed))
    kept_indices = [i for _, i in indexed]
    return [items[i] for i in kept_indices], kept_indices


def extract_docx_header_footer_texts(filename=None, binary=None):
    doc = Document(filename) if binary is None else Document(BytesIO(binary))
    texts = set()
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                normalized = re.sub(r"\s+", " ", paragraph.text).strip()
                if normalized:
                    texts.add(normalized)
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        normalized = re.sub(r"\s+", " ", cell.text).strip()
                        if normalized:
                            texts.add(normalized)
    return texts


def remove_header_footer_docx_sections(items, header_footer_texts):
    if not header_footer_texts:
        return items

    filtered = []
    for item in items:
        text = _item_text(item)
        normalized = re.sub(r"\s+", " ", text).strip() if isinstance(text, str) else ""
        if normalized and normalized in header_footer_texts:
            continue
        filtered.append(item)
    return filtered


def remove_header_footer_html_blob(blob):
    soup = BeautifulSoup(blob, "html.parser")
    for element in soup.find_all(
        lambda tag: tag.name in {"header", "footer"}
        or tag.get("role") in {"banner", "contentinfo"}
    ):
        element.decompose()
    return str(soup).encode("utf-8")


def extract_word_outlines(filename, binary=None):
    doc = Document(filename) if binary is None else Document(BytesIO(binary))
    outlines = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        match = re.search(r"Heading\s*(\d+)", style_name, re.I)
        if not match:
            continue
        outlines.append((text, int(match.group(1)) - 1, None))
    return outlines


def remove_toc_pdf(items, outlines):
    if not outlines:
        return items

    toc_start_page = None
    content_start_page = None
    for i, (title, level, page_no) in enumerate(outlines):
        if re.match(r"(contents|目录|目次|table of contents|致谢|acknowledge)$", title.split("@@")[0].strip().lower()):
            toc_start_page = page_no
            for next_title, next_level, next_page_no in outlines[i + 1:]:
                if next_level != level:
                    continue
                if re.match(r"(contents|目录|目次|table of contents|致谢|acknowledge)$", next_title.split("@@")[0].strip().lower()):
                    continue
                content_start_page = next_page_no
                break
            break

    if content_start_page:
        return [item for item in items if not (toc_start_page <= item["page_number"] < content_start_page)]
    return items


def remove_toc_word(items, outlines):
    if not outlines:
        filtered_items, _ = remove_toc(items)
        return filtered_items
    outline_titles = [title.split("@@")[0].strip().lower() for title, _, _ in outlines if title]
    if outline_titles:
        indexed = [(_item_text(item), i) for i, item in enumerate(items)]
        i = 0
        while i < len(indexed):
            if not re.match(r"(contents|目录|目次|table of contents|致谢|acknowledge)$", indexed[i][0].split("@@")[0].strip().lower()):
                i += 1
                continue
            indexed.pop(i)
            while i < len(indexed):
                text = indexed[i][0]
                normalized = text.split("@@")[0].strip().lower()
                if not normalized:
                    indexed.pop(i)
                    continue
                if any(normalized.startswith(title) or title.startswith(normalized) for title in outline_titles):
                    indexed.pop(i)
                    continue
                if re.search(r"(\.{2,}|…{2,}|·{2,}|[ ]{2,})\s*\d+\s*$", text):
                    indexed.pop(i)
                    continue
                break
            break
        items = [items[i] for _, i in indexed]
    filtered_items, _ = remove_toc(items)
    return filtered_items


def _item_text(item):
    if isinstance(item, str):
        return item
    if isinstance(item, dict):
        return item["text"]
    return item[0]


def _is_english(indexed):
    texts = [text for text, _ in indexed if text]
    if not texts:
        return False
    return is_english(random_choices(texts, k=200))


def enhance_media_sections_with_vision(
    sections,
    tenant_id,
    vlm_conf=None,
    callback=None,
):
    if not sections or not tenant_id:
        return sections

    try:
        try:
            vision_model_config = get_model_config_by_type_and_name(
                tenant_id, LLMType.IMAGE2TEXT, vlm_conf["llm_id"]
            )
        except Exception:
            vision_model_config = get_tenant_default_model_by_type(
                tenant_id, LLMType.IMAGE2TEXT
            )
        vision_model = LLMBundle(tenant_id, vision_model_config)
    except Exception:
        return sections

    for item in sections:
        if item.get("doc_type_kwd") not in {"image", "table"}:
            continue
        if item.get("image") is None:
            continue
        
        text = item.get("text") or ""
        try:
            parsed = VisionFigureParser(
                vision_model=vision_model,
                figures_data=[((item["image"], [""]), [(0, 0, 0, 0, 0)])],
                context_size=0,
            )(callback=callback)
        except Exception:
            continue

        if not parsed:
            continue

        # VisionFigureParser returns [((image, text_or_text_list), positions), ...].
        first_result = parsed[0]
        # first_result[0] is the (image, parsed_text) tuple.
        image_and_text = first_result[0]
        # image_and_text[1] is the parsed text content.
        parsed_text = str(image_and_text[1] or "").strip()

        if parsed_text:
            item["text"] = f"{text}\n{parsed_text}" if text else parsed_text

    return sections
