#
#  Copyright 2026 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from deepdoc.parser.docx_numbering import NumberedHeading, apply_numbered_headings_to_markdown, extract_numbered_headings


def _element(tag, value=None, **attributes):
    element = OxmlElement(tag)
    if value is not None:
        element.set(qn("w:val"), str(value))
    for name, attribute_value in attributes.items():
        element.set(qn(f"w:{name}"), str(attribute_value))
    return element


def _add_numbering_level(abstract, level, style_id, level_text, restart_level=None):
    definition = _element("w:lvl", ilvl=level)
    definition.append(_element("w:start", 1))
    definition.append(_element("w:numFmt", "decimal"))
    definition.append(_element("w:pStyle", style_id))
    definition.append(_element("w:lvlText", level_text))
    if restart_level is not None:
        definition.append(_element("w:lvlRestart", restart_level))
    abstract.append(definition)


def _set_style_numbering(style, number_id, level):
    properties = style.element.get_or_add_pPr()
    for existing in properties.findall(qn("w:numPr")):
        properties.remove(existing)
    numbering = _element("w:numPr")
    numbering.append(_element("w:ilvl", level))
    numbering.append(_element("w:numId", number_id))
    properties.append(numbering)


def _disable_paragraph_numbering(paragraph):
    properties = paragraph._p.get_or_add_pPr()
    numbering = _element("w:numPr")
    numbering.append(_element("w:ilvl", 0))
    numbering.append(_element("w:numId", 0))
    properties.append(numbering)


def _set_paragraph_numbering(paragraph, number_id, level):
    properties = paragraph._p.get_or_add_pPr()
    numbering = _element("w:numPr")
    numbering.append(_element("w:ilvl", level))
    numbering.append(_element("w:numId", number_id))
    properties.append(numbering)


def _numbered_document(restart_level=None, include_override_heading=False):
    document = Document()
    numbering_root = document.part.numbering_part.element
    abstract_id = 4242
    number_id = 4242

    abstract = _element("w:abstractNum", abstractNumId=abstract_id)
    _add_numbering_level(abstract, 0, "Heading1", "%1")
    _add_numbering_level(abstract, 1, "Heading2", "%1.%2", restart_level=restart_level)
    _add_numbering_level(abstract, 9, "Heading1", "%1")
    numbering_root.append(abstract)

    number = _element("w:num", numId=number_id)
    number.append(_element("w:abstractNumId", abstract_id))
    numbering_root.append(number)
    override_number_id = number_id + 1
    if include_override_heading:
        override_number = _element("w:num", numId=override_number_id)
        override_number.append(_element("w:abstractNumId", abstract_id))
        level_override = _element("w:lvlOverride", ilvl=0)
        level_override.append(_element("w:startOverride", 10))
        override_number.append(level_override)
        numbering_root.append(override_number)

    _set_style_numbering(document.styles["Heading 1"], number_id, 0)
    _set_style_numbering(document.styles["Heading 2"], number_id, 1)
    unnumbered_heading = document.styles.add_style("Unnumbered Heading", WD_STYLE_TYPE.PARAGRAPH)
    unnumbered_heading.base_style = document.styles["Heading 1"]
    _set_style_numbering(unnumbered_heading, 0, 0)
    document.add_paragraph("Аннотация", style=unnumbered_heading)
    introduction = document.add_paragraph("Введение", style="Heading 1")
    _disable_paragraph_numbering(introduction)
    document.add_paragraph("Общие сведения", style="Heading 1")
    body_list_item = document.add_paragraph("Обычный пункт списка")
    _set_paragraph_numbering(body_list_item, number_id, 0)
    outline_body = document.styles.add_style("Outline Body", WD_STYLE_TYPE.PARAGRAPH)
    outline_body_properties = outline_body.element.get_or_add_pPr()
    outline_body_properties.append(_element("w:outlineLvl", 0))
    _set_style_numbering(outline_body, number_id, 0)
    document.add_paragraph("Текст с outlineLvl", style=outline_body)
    document.add_paragraph("Порядок работы", style="Heading 1")
    document.add_paragraph("Начало работы", style="Heading 2")
    document.add_paragraph("Обслуживание", style="Heading 1")
    document.add_paragraph("Диагностика", style="Heading 2")
    if include_override_heading:
        override_heading = document.add_paragraph("Приложение", style="Heading 1")
        _set_paragraph_numbering(override_heading, override_number_id, 0)

    invalid_level = document.add_paragraph("Invalid level", style="Heading 1")
    invalid_properties = invalid_level._p.get_or_add_pPr()
    invalid_numbering = _element("w:numPr")
    invalid_numbering.append(_element("w:ilvl", 9))
    invalid_numbering.append(_element("w:numId", number_id))
    invalid_properties.append(invalid_numbering)

    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return Document(stream)


def test_extract_numbered_headings_from_style_numbering():
    headings = extract_numbered_headings(_numbered_document(), enabled=True)

    assert [(heading.numbered_text, heading.level) for heading in headings] == [
        ("1 Общие сведения", 1),
        ("2 Порядок работы", 1),
        ("2.1 Начало работы", 2),
        ("3 Обслуживание", 1),
        ("3.1 Диагностика", 2),
    ]


def test_explicit_numbering_disable_wins_over_numbered_style():
    headings = extract_numbered_headings(_numbered_document(), enabled=True)

    assert not any(heading.text in {"Аннотация", "Введение", "Обычный пункт списка", "Текст с outlineLvl"} for heading in headings)


def test_numbering_extraction_is_opt_in():
    assert extract_numbered_headings(_numbered_document()) == []


def test_numbering_extraction_can_be_disabled():
    headings = extract_numbered_headings(_numbered_document(), enabled=False)

    assert headings == []


def test_apply_numbered_headings_to_markdown():
    headings = extract_numbered_headings(_numbered_document(), enabled=True)

    markdown = "# Введение\n\n# Общие сведения\n\n# Порядок работы\n\n## Начало работы\n"

    assert apply_numbered_headings_to_markdown(markdown, headings) == "# Введение\n\n# 1 Общие сведения\n\n# 2 Порядок работы\n\n## 2.1 Начало работы"


def test_setext_detection_ignores_empty_list_and_table_candidates():
    for candidate in ("", "- List item", "1. List item", "Heading | Value"):
        markdown = f"{candidate}\n---"
        headings = [NumberedHeading(text=candidate, numbered_text=f"1 {candidate}", level=1)]

        assert apply_numbered_headings_to_markdown(markdown, headings) == markdown


def test_setext_detection_rejects_mixed_underline_markers():
    markdown = "Title\n=-"
    headings = [NumberedHeading(text="Title", numbered_text="1 Title", level=1)]

    assert apply_numbered_headings_to_markdown(markdown, headings) == markdown


def test_apply_numbered_headings_to_setext_markdown():
    headings = extract_numbered_headings(_numbered_document(), enabled=True)

    markdown = "Введение\n========\n\nОбщие сведения\n===============\n\nПорядок работы\n===============\n\nНачало работы\n-------------\n"

    assert apply_numbered_headings_to_markdown(markdown, headings) == "Введение\n========\n\n1 Общие сведения\n===============\n\n2 Порядок работы\n===============\n\n2.1 Начало работы\n-------------"


def test_lvl_restart_zero_preserves_lower_level_counter():
    headings = extract_numbered_headings(_numbered_document(restart_level=0), enabled=True)

    assert [heading.numbered_text for heading in headings][-1] == "3.2 Диагностика"


def test_new_numbering_instance_honors_start_override():
    headings = extract_numbered_headings(_numbered_document(include_override_heading=True), enabled=True)

    assert headings[-1].numbered_text == "10 Приложение"
