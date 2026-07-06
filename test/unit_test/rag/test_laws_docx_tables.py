#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import sys
import types
from importlib import import_module, reload
from io import BytesIO

import pytest
from docx import Document


def _stub(name, **attrs):
    mod = types.ModuleType(name)
    for key, value in attrs.items():
        setattr(mod, key, value)
    sys.modules.setdefault(name, mod)
    return mod


# Stub laws.py's app-layer siblings that the Docx parser never calls, so the module
# can be imported without pulling in the LLM / vision / storage stacks.
class _DummyBase:
    def __init__(self, *a, **k):
        pass


@pytest.fixture(scope="module")
def docx_chunker():
    original_modules = {
        name: sys.modules.get(name)
        for name in (
            "deepdoc.parser",
            "deepdoc.parser.utils",
            "rag.app.naive",
            "common.parser_config_utils",
        )
    }

    try:
        _stub("deepdoc.parser", PdfParser=_DummyBase, DocxParser=_DummyBase, HtmlParser=_DummyBase)
        _stub("deepdoc.parser.utils", get_text=lambda *a, **k: "")
        _stub("rag.app.naive", by_plaintext=lambda *a, **k: ([], [], None), PARSERS={})
        _stub("common.parser_config_utils", normalize_layout_recognizer=lambda x: (x, None))
        module = import_module("rag.app.laws")
        module = reload(module)
        yield module.Docx
    finally:
        for name, original in original_modules.items():
            if original is None:
                sys.modules.pop(name, None)
            else:
                sys.modules[name] = original


def _build_docx(builder):
    doc = Document()
    builder(doc)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.mark.p2
def test_laws_docx_preserves_table(docx_chunker):
    """Regression for #16008: the laws DOCX parser dropped tables entirely."""

    def builder(d):
        d.add_heading("Chapter 1 General Provisions", level=1)
        d.add_heading("Article 2 Fee Schedule", level=2)
        d.add_paragraph("The applicable fees are as follows:")
        t = d.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "Item"
        t.cell(0, 1).text = "Fee"
        t.cell(1, 0).text = "Registration"
        t.cell(1, 1).text = "100"

    chunks = docx_chunker()("law.docx", _build_docx(builder))

    assert any("<table>" in c for c in chunks)
    table_chunk = next(c for c in chunks if "<table>" in c)
    # Table content is present...
    assert "Registration" in table_chunk and "100" in table_chunk
    # ...and it carries its enclosing section's title path for retrieval context.
    assert "Article 2 Fee Schedule" in table_chunk


@pytest.mark.p2
def test_laws_docx_merged_cells_use_colspan(docx_chunker):
    def builder(d):
        d.add_heading("Heading", level=1)
        t = d.add_table(rows=1, cols=3)
        # Identical adjacent cell text is collapsed into a single colspan cell.
        t.cell(0, 0).text = "Merged"
        t.cell(0, 1).text = "Merged"
        t.cell(0, 2).text = "Other"

    chunks = docx_chunker()("law.docx", _build_docx(builder))
    table_chunk = next(c for c in chunks if "<table>" in c)
    assert "colspan='2'" in table_chunk
    assert "<td>Other</td>" in table_chunk


@pytest.mark.p2
def test_laws_docx_escapes_cell_html(docx_chunker):
    def builder(d):
        d.add_heading("Heading", level=1)
        t = d.add_table(rows=1, cols=1)
        t.cell(0, 0).text = "a < b & c > d"

    chunks = docx_chunker()("law.docx", _build_docx(builder))
    table_chunk = next(c for c in chunks if "<table>" in c)
    # Special characters are HTML-escaped so the table markup stays well-formed.
    assert "a &lt; b &amp; c &gt; d" in table_chunk
    assert "<td>a < b" not in table_chunk


@pytest.mark.p2
def test_laws_docx_tables_only_does_not_crash(docx_chunker):
    def builder(d):
        t = d.add_table(rows=1, cols=2)
        t.cell(0, 0).text = "a"
        t.cell(0, 1).text = "b"

    chunks = docx_chunker()("law.docx", _build_docx(builder))
    assert any("<table>" in c for c in chunks)


@pytest.mark.p2
def test_laws_docx_empty_doc_returns_empty(docx_chunker):
    chunks = docx_chunker()("law.docx", _build_docx(lambda d: None))
    assert chunks == []
