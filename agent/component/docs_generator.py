import logging
import base64
import json
import os
import re
import shutil
import tempfile
from abc import ABC
from datetime import datetime
from functools import partial
from io import BytesIO
from xml.sax.saxutils import escape

from agent.component.base import ComponentParamBase
from api.utils.api_utils import timeout
from .message import Message


def sanitize_filename(name: str, extension: str) -> str:
    if not name:
        return f"file.{extension}"

    name = str(name).strip()
    name = re.sub(r'[\\/\x00-\x1f\?\#\%\*\:\|\<\>"]', " ", name)
    name = re.sub(r"\s+", " ", name).strip(" .")

    if not name:
        return f"file.{extension}"

    base, _ = os.path.splitext(name)
    base = base[:180].rstrip() or "file"
    return f"{base}.{extension}"


class DocGeneratorParam(ComponentParamBase):
    """
    Define the Docs Generator component parameters.
    """

    def __init__(self):
        super().__init__()
        self.output_format = "pdf"  # pdf, docx, txt, markdown, html
        self.content = ""
        self.filename = ""
        self.header_text = ""
        self.footer_text = ""
        self.watermark_text = ""
        self.add_page_numbers = True
        self.add_timestamp = True
        self.font_size = 12
        self.outputs = {
            "download": {"value": "", "type": "string"},
        }

    def check(self):
        self.check_empty(self.content, "[DocGenerator] Content")
        self.check_valid_value(
            self.output_format,
            "[DocGenerator] Output format",
            ["pdf", "docx", "txt", "markdown", "html"],
        )
        self.check_positive_number(self.font_size, "[DocGenerator] Font size")
        if self.font_size < 12:
            raise ValueError("[DocGenerator] Font size must be greater than or equal to 12")


class DocGenerator(Message, ABC):
    component_name = "DocGenerator"
    _default_output_directory = os.path.join(tempfile.gettempdir(), "doc_outputs")
    _overlay_margin = 36
    _overlay_font_size = 9
    _pdf_main_font = "Noto Sans CJK SC"
    _pdf_cjk_font = "Noto Sans CJK SC"
    _pdf_overlay_font = "STSong-Light"

    def get_input_form(self) -> dict[str, dict]:
        return {
            "content": {
                "name": "Content",
                "type": "text",
            }
        }

    @timeout(int(os.environ.get("COMPONENT_EXEC_TIMEOUT", 10 * 60)))
    def _invoke(self, **kwargs):
        try:
            content = self._resolve_content(kwargs)
            output_format = self._param.output_format or "pdf"

            try:
                if output_format == "pdf":
                    file_path, doc_base64 = self._generate_pdf(content)
                    mime_type = "application/pdf"
                elif output_format == "docx":
                    file_path, doc_base64 = self._generate_docx(content)
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                elif output_format == "txt":
                    file_path, doc_base64 = self._generate_txt(content)
                    mime_type = "text/plain"
                elif output_format == "markdown":
                    file_path, doc_base64 = self._generate_markdown(content)
                    mime_type = "text/markdown"
                elif output_format == "html":
                    file_path, doc_base64 = self._generate_html(content)
                    mime_type = "text/html"
                else:
                    raise Exception(f"Unsupported output format: {output_format}")

                filename = os.path.basename(file_path)
                if not os.path.exists(file_path):
                    raise Exception("Document file was not created")

                file_size = os.path.getsize(file_path)
                if file_size == 0:
                    raise Exception("Document file is empty")

                logging.info(
                    "Successfully generated %s: %s (Size: %s bytes)",
                    output_format.upper(),
                    filename,
                    file_size,
                )

                download_info = {
                    "filename": filename,
                    "base64": doc_base64,
                    "mime_type": mime_type,
                    "size": file_size,
                }
                self.set_output("download", json.dumps(download_info))
                return download_info

            except Exception as e:
                logging.exception("Error generating %s document", output_format)
                self.set_output("_ERROR", f"Document generation failed: {str(e)}")
                raise

        except Exception as e:
            logging.exception("Error in DocGenerator._invoke")
            self.set_output("_ERROR", f"Document generation failed: {str(e)}")
            raise

    def _resolve_content(self, kwargs: dict) -> str:
        content = self._param.content or ""
        logging.info("Starting document generation, content length: %s chars", len(content))

        if content and self._canvas.is_reff(content.strip()):
            matches = re.findall(self.variable_ref_patt, content, flags=re.DOTALL)
            for match in matches:
                try:
                    var_value = self._canvas.get_variable_value(match)
                    if var_value is None:
                        continue
                    if isinstance(var_value, partial):
                        resolved_content = ""
                        for chunk in var_value():
                            resolved_content += chunk
                        content = content.replace("{" + match + "}", resolved_content)
                    else:
                        content = content.replace("{" + match + "}", str(var_value))
                except Exception as e:
                    logging.warning("Error resolving variable %s: %s", match, str(e))
                    content = content.replace("{" + match + "}", f"[ERROR: {str(e)}]")

        if content:
            try:
                content, _ = self.get_kwargs(content, kwargs)
            except Exception as e:
                logging.warning("Error processing content with get_kwargs: %s", str(e))

        if not content:
            content = kwargs.get("content", "")

        return content

    def _get_output_directory(self) -> str:
        os.makedirs(self._default_output_directory, exist_ok=True)
        return self._default_output_directory

    def _build_output_filename(self, output_format: str) -> str:
        import uuid

        if self._param.filename:
            return sanitize_filename(self._param.filename, output_format.lower())

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"document_{timestamp}_{uuid.uuid4().hex[:8]}.{output_format}"

    def _get_timestamp_text(self) -> str:
        return f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"

    def _write_bytes_output(self, content: bytes, extension: str) -> tuple[str, str]:
        output_directory = self._get_output_directory()
        filename = self._build_output_filename(extension)
        file_path = os.path.join(output_directory, filename)
        with open(file_path, "wb") as f:
            f.write(content)
        return file_path, base64.b64encode(content).decode("utf-8")

    def _build_markdown_source(self, content: str, include_timestamp_in_body: bool = False) -> str:
        if not (include_timestamp_in_body and self._param.add_timestamp):
            return content
        return f"{self._get_timestamp_text()}\n\n{content}"

    def _get_heading_sizes(self) -> tuple[int, int, int]:
        base = int(self._param.font_size)
        return base + 6, base + 4, base + 2

    def _generate_pandoc_binary_output(
        self,
        content: str,
        target_format: str,
        extension: str,
        include_timestamp_in_body: bool = False,
        extra_args: list[str] | None = None,
    ) -> tuple[str, str]:
        import pypandoc

        output_directory = self._get_output_directory()
        filename = self._build_output_filename(extension)
        file_path = os.path.join(output_directory, filename)
        markdown_content = self._build_markdown_source(
            content,
            include_timestamp_in_body=include_timestamp_in_body,
        )

        pypandoc.convert_text(
            markdown_content,
            to=target_format,
            format="markdown",
            outputfile=file_path,
            extra_args=extra_args or [],
        )

        with open(file_path, "rb") as f:
            file_bytes = f.read()

        return file_path, base64.b64encode(file_bytes).decode("utf-8")

    def _generate_pandoc_text_output(
        self,
        content: str,
        target_format: str,
        extension: str,
        include_timestamp_in_body: bool = True,
    ) -> tuple[str, str]:
        import pypandoc

        markdown_content = self._build_markdown_source(
            content,
            include_timestamp_in_body=include_timestamp_in_body,
        )
        converted_content = pypandoc.convert_text(
            markdown_content,
            to=target_format,
            format="markdown",
        )
        return self._write_bytes_output(converted_content.encode("utf-8"), extension)

    def _select_pdf_engine(self) -> str:
        if shutil.which("xelatex"):
            return "xelatex"
        raise Exception("No PDF engine found. Install xelatex.")

    def _get_pdf_font_args(self) -> list[str]:
        return [
            "-V",
            f"mainfont={self._pdf_main_font}",
            "-V",
            f"CJKmainfont={self._pdf_cjk_font}",
        ]

    def _get_pdf_overlay_font_name(self) -> str:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont

        try:
            pdfmetrics.getFont(self._pdf_overlay_font)
        except KeyError:
            pdfmetrics.registerFont(UnicodeCIDFont(self._pdf_overlay_font))

        return self._pdf_overlay_font

    def _build_pdf_heading_overrides(self) -> str:
        font_size = int(self._param.font_size)
        leading = round(font_size * 1.2, 1)
        h1_size, h2_size, h3_size = self._get_heading_sizes()
        h1_leading = round(h1_size * 1.2, 1)
        h2_leading = round(h2_size * 1.2, 1)
        h3_leading = round(h3_size * 1.2, 1)

        return rf"""
\makeatletter
\renewcommand\normalsize{{
  \@setfontsize\normalsize{{{font_size}pt}}{{{leading}pt}}
  \abovedisplayskip 12pt plus 3pt minus 7pt
  \abovedisplayshortskip \z@ plus 3pt
  \belowdisplayshortskip 6.5pt plus 3.5pt minus 3pt
  \belowdisplayskip \abovedisplayskip
  \let\@listi\@listI
}}
\normalsize
\renewcommand\section{{\@startsection{{section}}{{1}}{{\z@}}{{-3.5ex \@plus -1ex \@minus -.2ex}}{{2.3ex \@plus .2ex}}{{\normalfont\fontsize{{{h1_size}pt}}{{{h1_leading}pt}}\selectfont\bfseries}}}}
\renewcommand\subsection{{\@startsection{{subsection}}{{2}}{{\z@}}{{-3.25ex\@plus -1ex \@minus -.2ex}}{{1.5ex \@plus .2ex}}{{\normalfont\fontsize{{{h2_size}pt}}{{{h2_leading}pt}}\selectfont\bfseries}}}}
\renewcommand\subsubsection{{\@startsection{{subsubsection}}{{3}}{{\z@}}{{-3.25ex\@plus -1ex \@minus -.2ex}}{{1.5ex \@plus .2ex}}{{\normalfont\fontsize{{{h3_size}pt}}{{{h3_leading}pt}}\selectfont\bfseries}}}}
\makeatother
""".strip()

    def _write_temp_tex(self, content: str) -> str:
        output_directory = self._get_output_directory()
        with tempfile.NamedTemporaryFile(
            mode="w",
            encoding="utf-8",
            suffix=".tex",
            dir=output_directory,
            delete=False,
        ) as f:
            f.write(content)
            return f.name

    def _should_apply_pdf_overlay(self) -> bool:
        return any(
            [
                self._param.header_text,
                self._param.footer_text,
                self._param.watermark_text,
                self._param.add_page_numbers,
                self._param.add_timestamp,
            ]
        )

    def _build_pdf_overlay_page(self, width: float, height: float, page_number: int):
        if not self._should_apply_pdf_overlay():
            return None

        from pypdf import PdfReader
        from reportlab.lib.colors import Color
        from reportlab.pdfgen import canvas as pdf_canvas

        buffer = BytesIO()
        overlay = pdf_canvas.Canvas(buffer, pagesize=(width, height))
        overlay_font = self._get_pdf_overlay_font_name()

        if self._param.watermark_text:
            overlay.saveState()
            if hasattr(overlay, "setFillAlpha"):
                overlay.setFillAlpha(0.15)
            overlay.setFillColor(Color(0.6, 0.6, 0.6))
            overlay.setFont(overlay_font, 48)
            overlay.translate(width / 2, height / 2)
            overlay.rotate(45)
            overlay.drawCentredString(0, 0, self._param.watermark_text)
            overlay.restoreState()

        overlay.setFont(overlay_font, self._overlay_font_size)
        overlay.setFillColor(Color(0.35, 0.35, 0.35))

        if self._param.header_text:
            overlay.drawString(
                self._overlay_margin,
                height - self._overlay_margin + 8,
                self._param.header_text,
            )

        if self._param.footer_text:
            overlay.drawString(
                self._overlay_margin,
                self._overlay_margin - 8,
                self._param.footer_text,
            )

        if self._param.add_timestamp:
            overlay.drawCentredString(
                width / 2,
                self._overlay_margin - 8,
                self._get_timestamp_text(),
            )

        if self._param.add_page_numbers:
            overlay.drawRightString(
                width - self._overlay_margin,
                self._overlay_margin - 8,
                f"Page {page_number}",
            )

        overlay.save()
        buffer.seek(0)
        return PdfReader(buffer).pages[0]

    def _apply_pdf_overlay(self, file_path: str) -> tuple[str, str]:
        from pypdf import PdfReader, PdfWriter

        if not self._should_apply_pdf_overlay():
            with open(file_path, "rb") as f:
                file_bytes = f.read()
            return file_path, base64.b64encode(file_bytes).decode("utf-8")

        reader = PdfReader(file_path)
        writer = PdfWriter()

        for page_number, page in enumerate(reader.pages, start=1):
            overlay_page = self._build_pdf_overlay_page(
                float(page.mediabox.width),
                float(page.mediabox.height),
                page_number,
            )
            if overlay_page is not None:
                page.merge_page(overlay_page)
            writer.add_page(page)

        temp_file = f"{file_path}.overlay"
        with open(temp_file, "wb") as f:
            writer.write(f)

        os.replace(temp_file, file_path)
        with open(file_path, "rb") as f:
            file_bytes = f.read()
        return file_path, base64.b64encode(file_bytes).decode("utf-8")

    def _clear_docx_container(self, container):
        element = container._element
        for child in list(element):
            element.remove(child)

    def _append_docx_field(self, run, instruction: str):
        from docx.oxml import OxmlElement

        begin = OxmlElement("w:fldChar")
        begin.set(run.part.element.nsmap["w"] and "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType", "begin")

        instr = OxmlElement("w:instrText")
        instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        instr.text = instruction

        end = OxmlElement("w:fldChar")
        end.set(run.part.element.nsmap["w"] and "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType", "end")

        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    def _add_docx_watermark(self, section):
        if not self._param.watermark_text:
            return

        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml

        header = section.header
        paragraph = header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        watermark_xml = parse_xml(
            rf"""
            <w:pict
              xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
              xmlns:v="urn:schemas-microsoft-com:vml"
              xmlns:o="urn:schemas-microsoft-com:office:office">
              <v:shape id="PowerPlusWaterMarkObject"
                o:spid="_x0000_s2049"
                type="#_x0000_t136"
                style="position:absolute;
                  margin-left:0;
                  margin-top:0;
                  width:468pt;
                  height:117pt;
                  rotation:315;
                  z-index:-251654144;
                  mso-wrap-edited:f;
                  mso-position-horizontal:center;
                  mso-position-horizontal-relative:margin;
                  mso-position-vertical:center;
                  mso-position-vertical-relative:margin"
                fillcolor="#d9d9d9"
                stroked="f">
                <v:fill opacity="0.18"/>
                <v:textpath on="t" style="font-family:&quot;Calibri&quot;;font-size:1pt" string="{escape(self._param.watermark_text)}"/>
              </v:shape>
            </w:pict>
            """
        )
        run._r.append(watermark_xml)

    def _normalize_docx_section_geometry(self, section, default_section):
        for attr in ("page_width", "left_margin", "right_margin"):
            if getattr(section, attr) is None:
                setattr(section, attr, getattr(default_section, attr))

    def _get_docx_available_width(self, section):
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin

        if page_width is None or left_margin is None or right_margin is None:
            raise ValueError("DOCX section geometry is incomplete after normalization.")

        return page_width - left_margin - right_margin

    def _decorate_docx(self, file_path: str) -> tuple[str, str]:
        from docx import Document
        from docx.enum.text import WD_TAB_ALIGNMENT
        from docx.shared import Pt

        document = Document(file_path)
        default_section = Document().sections[0]
        h1_size, h2_size, h3_size = self._get_heading_sizes()

        style_map = {
            "Normal": int(self._param.font_size),
            "Heading 1": h1_size,
            "Heading 2": h2_size,
            "Heading 3": h3_size,
        }
        for style_name, size in style_map.items():
            try:
                document.styles[style_name].font.size = Pt(size)
            except Exception:
                continue

        for section in document.sections:
            self._normalize_docx_section_geometry(section, default_section)
            available_width = self._get_docx_available_width(section)

            header = section.header
            header.is_linked_to_previous = False
            self._clear_docx_container(header)
            if self._param.header_text:
                paragraph = header.add_paragraph()
                paragraph.add_run(self._param.header_text)

            self._add_docx_watermark(section)

            footer = section.footer
            footer.is_linked_to_previous = False
            self._clear_docx_container(footer)
            if any(
                [
                    self._param.footer_text,
                    self._param.add_timestamp,
                    self._param.add_page_numbers,
                ]
            ):
                paragraph = footer.add_paragraph()
                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    int(available_width // 2),
                    WD_TAB_ALIGNMENT.CENTER,
                )
                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    int(available_width),
                    WD_TAB_ALIGNMENT.RIGHT,
                )

                if self._param.footer_text:
                    paragraph.add_run(self._param.footer_text)

                if self._param.add_timestamp or self._param.add_page_numbers:
                    paragraph.add_run("\t")

                if self._param.add_timestamp:
                    paragraph.add_run(self._get_timestamp_text())

                if self._param.add_page_numbers:
                    paragraph.add_run("\t")
                    self._append_docx_field(paragraph.add_run(), " PAGE ")

        document.save(file_path)
        with open(file_path, "rb") as f:
            file_bytes = f.read()
        return file_path, base64.b64encode(file_bytes).decode("utf-8")

    def thoughts(self) -> str:
        return f"Generating {self._param.output_format.upper()} document with markdown conversion..."

    def _generate_pdf(self, content: str) -> tuple[str, str]:
        try:
            engine = self._select_pdf_engine()
            header_path = self._write_temp_tex(self._build_pdf_heading_overrides())
            try:
                file_path, _ = self._generate_pandoc_binary_output(
                    content,
                    "pdf",
                    "pdf",
                    include_timestamp_in_body=False,
                    extra_args=[
                        "--standalone",
                        f"--pdf-engine={engine}",
                        f"--include-in-header={header_path}",
                        *self._get_pdf_font_args(),
                    ],
                )
            finally:
                if os.path.exists(header_path):
                    os.remove(header_path)
            return self._apply_pdf_overlay(file_path)
        except Exception as e:
            raise Exception(f"PDF generation failed: {str(e)}")

    def _generate_docx(self, content: str) -> tuple[str, str]:
        try:
            file_path, _ = self._generate_pandoc_binary_output(
                content,
                "docx",
                "docx",
                include_timestamp_in_body=False,
                extra_args=["--standalone"],
            )
            return self._decorate_docx(file_path)
        except Exception as e:
            raise Exception(f"DOCX generation failed: {str(e)}")

    def _generate_txt(self, content: str) -> tuple[str, str]:
        try:
            return self._generate_pandoc_text_output(content, "plain", "txt")
        except Exception as e:
            raise Exception(f"TXT generation failed: {str(e)}")

    def _generate_markdown(self, content: str) -> tuple[str, str]:
        try:
            return self._generate_pandoc_text_output(content, "markdown", "md")
        except Exception as e:
            raise Exception(f"Markdown generation failed: {str(e)}")

    def _generate_html(self, content: str) -> tuple[str, str]:
        try:
            return self._generate_pandoc_text_output(content, "html", "html")
        except Exception as e:
            raise Exception(f"HTML generation failed: {str(e)}")
